# gcpl_pricing_hybrid_full.py
"""
GCPL Pricing Hybrid - Full Single File (UPDATED: NL safety + improved parsing + LLM summarization constraints)
Run:
    source venv/bin/activate
    python gcpl_pricing_hybrid_full.py
"""
import os
import json
import math
import time
import random
import argparse
import re
import difflib
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional, Dict, Any, List

import numpy as np
import pandas as pd

# UI
from rich.console import Console
from rich.table import Table
from rich.panel import Panel

console = Console(width=160)

# Feature toggles - set True if you will install packages
ENABLE_NN = True
ENABLE_RL = True

# Try import NN & RL libs; if absent will disable features with clear message
if ENABLE_NN:
    try:
        import torch
        import torch.nn as nn
        import torch.optim as optim
    except Exception as e:
        console.print(f"Torch missing or failed to import: {e}", markup=False)
        console.print("Set ENABLE_NN=False or install torch in your venv.", markup=False)
        ENABLE_NN = False

if ENABLE_RL:
    try:
        from sklearn.ensemble import RandomForestClassifier
        import joblib
    except Exception as e:
        console.print(f"sklearn/joblib missing or failed to import: {e}", markup=False)
        console.print("Set ENABLE_RL=False or install scikit-learn & joblib.", markup=False)
        ENABLE_RL = False

# docx
try:
    from docx import Document
except Exception:
    Document = None

# boto3 (Bedrock)
try:
    import boto3
except Exception:
    boto3 = None

try:
    from flask import Flask, jsonify, request
except Exception:
    Flask = None

# Paths
BASE_DIR = Path(__file__).resolve().parent

def _load_env_file(path: Path) -> None:
    """Load KEY=value lines from path into os.environ. File wins over existing env."""
    try:
        if not path.exists():
            return
        for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            if "=" in line:
                k, v = line.split("=", 1)
                k = k.strip()
                v = v.strip().strip("'\"").replace("\\n", "\n")
                if k and v:
                    os.environ[k] = v
    except Exception:
        pass

# Load .env / .env.local so AWS_* are available without exporting in terminal every time.
# Try dotenv first, then always run our loader from BASE_DIR so credentials are never missed.
try:
    from dotenv import load_dotenv
    load_dotenv(BASE_DIR / ".env")
    load_dotenv(BASE_DIR / ".env.local")
except ImportError:
    pass
for _env_path in (BASE_DIR / ".env", BASE_DIR / ".env.local", Path.cwd() / ".env", Path.cwd() / ".env.local"):
    _load_env_file(_env_path)
DATA_CSV = os.getenv("AI_SAHAYAK_DATA_CSV", str(BASE_DIR / "raju_kirana_2yr.csv"))
SOP_DOCX = os.getenv(
    "AI_SAHAYAK_SOP_DOCX",
    str(BASE_DIR / "Dynamic Pricing Simulator and Promo Optimizer Build Requirements.docx"),
)
OUTPUT_DIR = os.getenv("AI_SAHAYAK_OUTPUT_DIR", str(BASE_DIR / "gcpl_pricing_outputs"))
AWS_REGION = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
DEEPAR_ENDPOINT = os.getenv("AI_SAHAYAK_DEEPAR_ENDPOINT", "")
CALENDAR_EVENTS_JSON = os.getenv("AI_SAHAYAK_CALENDAR_EVENTS_JSON", "[]")

DATASET_CATALOG = {
    "raju": str(BASE_DIR / "raju_kirana_2yr.csv"),
    "ramesh": str(BASE_DIR / "ramesh_msme_2yr.csv"),
    "suresh": str(BASE_DIR / "suresh_msme_2yr.csv"),
    "kanta": str(BASE_DIR / "kanta_msme_2yr.csv"),
    "lakshmi": str(BASE_DIR / "lakshmi_msme_2yr.csv"),
}

RAW_READ_COLUMNS = {
    "date",
    "txn_timestamp",
    "sku_id",
    "item_name",
    "item_category",
    "quantity_sold",
    "units_sold",
    "selling_price",
    "local_price",
    "discount_pct",
    "promo_depth_pct",
    "effective_price",
    "final_price",
    "market_price",
    "competitor_price",
    "opening_stock",
    "stock_on_hand",
    "inventory_level",
    "reorder_point",
    "needs_reorder",
    "empty_shelf_flag",
    "is_stockout",
    "revenue",
    "cost_amt",
    "cost",
    "profit_amt",
    "gross_margin_pct",
    "margin",
    "festival_impact",
    "festival_lift",
    "festival_name",
    "is_festival_day",
    "is_pre_festival_window",
    "promo_flag",
    "payment_mode",
    "is_weekend",
    "season",
    "weather_tag",
    "lead_time_days",
    "purchase_cost",
    "mrp",
    "tax_pct",
    "supplier_id",
}

READ_DTYPES = {
    "sku_id": "string",
    "item_name": "string",
    "item_category": "string",
    "festival_name": "string",
    "payment_mode": "string",
    "season": "string",
    "weather_tag": "string",
    "supplier_id": "string",
    "store_city": "string",
}

# LLM model and token limit
BEDROCK_MODEL_ID = os.getenv("AI_SAHAYAK_BEDROCK_MODEL_ID", "amazon.nova-pro-v1:0")
BEDROCK_FALLBACK_MODELS = [
    m.strip() for m in os.getenv(
        "AI_SAHAYAK_BEDROCK_FALLBACKS",
        "amazon.nova-lite-v1:0,anthropic.claude-3-5-sonnet-20241022-v2:0,openai.gpt-oss-20b-1:0",
    ).split(",")
    if m.strip()
]
LLM_MAX_TOKENS = 1200
REQUIRE_BEDROCK = os.getenv("AI_SAHAYAK_REQUIRE_BEDROCK", "1").strip().lower() not in ("0", "false", "no")
_BEDROCK_STATUS_CACHE: Dict[str, Any] = {"checked_at": 0.0, "ok": False, "error": "not_checked"}

# Chat-driven UI: when WP user asks for review/price, agents POST here; dashboard frontend polls GET to auto-navigate.
_CHAT_DRIVE_STATE: Dict[str, Any] = {}
_CHAT_DRIVE_TTL_SEC = 30

# Default SOP guardrails
DEFAULT_GUARDRAILS = {
    "min_margin_pct": 12.0,
    "max_promo_depth_pct": 40.0,
    "auto_approve_pct": 2.0,
    "manager_review_pct": 5.0,
    "max_monthly_change_pct": 15.0
}

# ---------------- Utilities ----------------
def ensure_output_dir() -> str:
    Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR

def ensure_parent_dir(path: str) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)


def get_deepar_endpoint(dataset_key: Optional[str] = None) -> str:
    key = str(dataset_key or "").strip().upper()
    if key:
        specific = os.getenv(f"AI_SAHAYAK_DEEPAR_ENDPOINT_{key}", "").strip()
        if specific:
            return specific
    return DEEPAR_ENDPOINT.strip()

def safe_float(x, default=float("nan")) -> float:
    try:
        if pd.isna(x): return default
    except Exception:
        pass
    try:
        if x is None: return default
        return float(x)
    except Exception:
        return default

def pct_change(a, b) -> float:
    try:
        a_f, b_f = float(a), float(b)
    except Exception:
        return 0.0
    if a_f == 0.0:
        return 0.0 if b_f == 0.0 else 100.0
    return (b_f - a_f) / (abs(a_f) + 1e-9) * 100.0

def format_inr(x):
    try:
        v = float(x)
        return "₹{:,.2f}".format(v)
    except Exception:
        return str(x)

def safe_num_for_json(x, default=None):
    try:
        if x is None: return default
        if isinstance(x, (np.floating, float)):
            if np.isfinite(x):
                return float(x)
            return default
        if isinstance(x, (np.integer, int)):
            return int(x)
        if isinstance(x, dict):
            return {k: safe_num_for_json(v, None) for k, v in x.items()}
        if isinstance(x, list):
            return [safe_num_for_json(v, None) for v in x]
        return x
    except Exception:
        return default

def normalize_token(value: Any) -> str:
    if value is None:
        return ""
    s = str(value).strip().lower()
    s = re.sub(r"[^a-z0-9]+", "", s)
    return s

def resolve_input_value(raw_value: str, candidates: List[str], label: str, cutoff: float = 0.62) -> str:
    raw = (raw_value or "").strip()
    if not raw:
        return raw
    if not candidates:
        return raw

    # Exact case-sensitive match fast-path
    if raw in candidates:
        return raw

    # Normalized exact match
    norm_to_original: Dict[str, List[str]] = {}
    for c in candidates:
        key = normalize_token(c)
        norm_to_original.setdefault(key, []).append(c)
    raw_norm = normalize_token(raw)
    if raw_norm in norm_to_original:
        chosen = norm_to_original[raw_norm][0]
        if chosen != raw:
            console.print(f"{label} normalized: '{raw}' -> '{chosen}'", markup=False)
        return chosen

    # Substring contains fallback (for partial SKU fragments)
    partial = [c for c in candidates if raw_norm and raw_norm in normalize_token(c)]
    if len(partial) == 1:
        chosen = partial[0]
        console.print(f"{label} inferred from partial input: '{raw}' -> '{chosen}'", markup=False)
        return chosen

    # Fuzzy match for typos
    keys = list(norm_to_original.keys())
    guess = difflib.get_close_matches(raw_norm, keys, n=1, cutoff=cutoff)
    if guess:
        chosen = norm_to_original[guess[0]][0]
        console.print(f"{label} typo-corrected: '{raw}' -> '{chosen}'", markup=False)
        return chosen

    return raw

# ---------------- JSON helpers (make AgentTrace JSON-serializable) ----------------
def to_serializable_traces(traces: Dict[str, Any]) -> Dict[str, Any]:
    out = {}
    for k, v in traces.items():
        try:
            out[k] = {
                "name": getattr(v, "name", str(k)),
                "candidates": {lvl: safe_num_for_json(price, None) for lvl, price in getattr(v, "candidates", {}).items()},
                "details": {dk: safe_num_for_json(dv, None) for dk, dv in getattr(v, "details", {}).items()}
            }
        except Exception:
            try:
                if isinstance(v, dict):
                    out[k] = {kk: safe_num_for_json(vv, None) for kk, vv in v.items()}
                else:
                    out[k] = str(v)
            except Exception:
                out[k] = str(v)
    return out

def serialize_pipeline_result(result: Dict[str, Any]) -> Dict[str, Any]:
    out = {}
    for k, v in result.items():
        if k == "traces":
            out["traces"] = to_serializable_traces(v)
        elif k in ("candidates", "agent_explanations", "candidate_explanations"):
            try:
                if isinstance(v, dict):
                    out[k] = {kk: {k2: safe_num_for_json(v2, None) for k2, v2 in vv.items()} for kk, vv in v.items()}
                else:
                    out[k] = v
            except Exception:
                out[k] = v
        else:
            try:
                if isinstance(v, (dict, list, str, bool)) or v is None:
                    out[k] = v
                else:
                    out[k] = safe_num_for_json(v, v)
            except Exception:
                out[k] = str(v)
    return out

# ---------------- SOP docx writer/reader ----------------
def write_detailed_sop(path=SOP_DOCX):
    if Document is None:
        console.print("[yellow]python-docx not installed. Can't write SOP docx.[/yellow]")
        return
    if os.path.exists(path):
        console.print(f"SOP already exists at {path}", markup=False)
        return
    ensure_parent_dir(path)
    doc = Document()
    doc.add_heading("Dynamic Pricing Simulator and Promo Optimizer - SOP", level=1)
    doc.add_paragraph("Purpose: Provide guardrails, rules, and operational definitions for automated pricing decisions.")
    doc.add_heading("Guardrails", level=2)
    doc.add_paragraph("min_margin_pct: 12.0  ; Minimum acceptable margin percent")
    doc.add_paragraph("max_promo_depth_pct: 40.0  ; Maximum allowed promotion depth")
    doc.add_paragraph("auto_approve_pct: 2.0  ; Changes <= 2% auto-approved")
    doc.add_paragraph("manager_review_pct: 5.0  ; Changes <= 5% require manager review")
    doc.add_paragraph("max_monthly_change_pct: 15.0  ; Max monthly change allowed")
    doc.add_heading("Business Rules", level=2)
    doc.add_paragraph("- Pricing must not drop below min_margin_pct unless explicit override.\n- Promotions in ecommerce allowed up to max_promo_depth_pct.\n- Channel-specific rules:\n  * ECOM: promotions allowed\n  * GT: margin-preserving preferred\n  * MT: manager review preferred for changes > 3%")
    doc.add_heading("Inventory Rules", level=2)
    doc.add_paragraph("- inventory_days < 10: raise price by 3% (neutral), 6% (aggressive)\n- inventory_days > 60: reduce price by 3% (neutral), 6% (aggressive)")
    doc.add_heading("Forecasting", level=2)
    doc.add_paragraph("- Forecast uses monthly seasonality index computed from units_sold across historical data.\n- Forecast uses SKU-level trend where available, else category or global trend.")
    doc.add_heading("Notes", level=2)
    doc.add_paragraph("This is a generated SOP used for testing and should be replaced by official SOP in production.")
    doc.save(path)
    console.print(f"[green]Wrote detailed SOP to {path}[/green]")

def parse_sop_docx(path=SOP_DOCX):
    guardrails = DEFAULT_GUARDRAILS.copy()
    raw = ""
    if not os.path.exists(path):
        write_detailed_sop(path)
    if not os.path.exists(path) or Document is None:
        console.print("SOP docx unreadable or missing; using defaults.", markup=False)
        return {"guardrails": guardrails, "raw": raw, "source": None}
    try:
        doc = Document(path)
        paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        raw = "\n".join(paras)
        m = re.search(r"min_margin_pct[:\s]*([0-9]+(?:\.\d+)?)", raw, flags=re.I)
        if m: guardrails["min_margin_pct"] = float(m.group(1))
        m = re.search(r"max_promo_depth_pct[:\s]*([0-9]+(?:\.\d+)?)", raw, flags=re.I)
        if m: guardrails["max_promo_depth_pct"] = float(m.group(1))
        m = re.search(r"auto_approve_pct[:\s]*([0-9]+(?:\.\d+)?)", raw, flags=re.I)
        if m: guardrails["auto_approve_pct"] = float(m.group(1))
        m = re.search(r"manager_review_pct[:\s]*([0-9]+(?:\.\d+)?)", raw, flags=re.I)
        if m: guardrails["manager_review_pct"] = float(m.group(1))
        m = re.search(r"max_monthly_change_pct[:\s]*([0-9]+(?:\.\d+)?)", raw, flags=re.I)
        if m: guardrails["max_monthly_change_pct"] = float(m.group(1))
        console.print(f"SOP parsed. Guardrails: {guardrails} (source: {path})", markup=False)
        return {"guardrails": guardrails, "raw": raw, "source": path}
    except Exception as e:
        console.print(f"SOP parse failed ({e}). Using defaults.", markup=False)
        return {"guardrails": guardrails, "raw": raw, "source": path}

# ---------------- Data loader + seasonality enrichment ----------------
def load_and_enrich(path=DATA_CSV) -> pd.DataFrame:
    if not os.path.exists(path):
        console.print(f"Dataset not found at {path}. Generating a synthetic demo dataset...", markup=False)
        ensure_parent_dir(path)
        generate_synthetic_dataset(path)
    df = pd.read_csv(
        path,
        usecols=lambda c: c in RAW_READ_COLUMNS,
        parse_dates=["date"],
        dtype=READ_DTYPES,
        low_memory=True,
    )
    if "txn_timestamp" in df.columns:
        df["txn_timestamp"] = pd.to_datetime(df["txn_timestamp"], errors="coerce")
    # Schema normalization for richer kirana datasets.
    if "local_price" not in df.columns and "selling_price" in df.columns:
        df["local_price"] = pd.to_numeric(df["selling_price"], errors="coerce")
    if "final_price" not in df.columns:
        if "effective_price" in df.columns:
            df["final_price"] = pd.to_numeric(df["effective_price"], errors="coerce")
        elif "local_price" in df.columns:
            df["final_price"] = pd.to_numeric(df["local_price"], errors="coerce")
    if "units_sold" not in df.columns and "quantity_sold" in df.columns:
        df["units_sold"] = pd.to_numeric(df["quantity_sold"], errors="coerce")
    if "inventory_level" not in df.columns and "stock_on_hand" in df.columns:
        df["inventory_level"] = pd.to_numeric(df["stock_on_hand"], errors="coerce")
    if "is_stockout" not in df.columns and "empty_shelf_flag" in df.columns:
        df["is_stockout"] = pd.to_numeric(df["empty_shelf_flag"], errors="coerce")
    if "promo_depth_pct" not in df.columns and "discount_pct" in df.columns:
        df["promo_depth_pct"] = pd.to_numeric(df["discount_pct"], errors="coerce") * 100.0
    if "promo_flag" not in df.columns:
        df["promo_flag"] = (pd.to_numeric(df.get("promo_depth_pct", 0), errors="coerce").fillna(0) > 0).astype(int)
    if "category" not in df.columns and "item_category" in df.columns:
        df["category"] = df["item_category"].astype(str)
    if "channel" not in df.columns:
        df["channel"] = "GT"
    if "region" not in df.columns:
        df["region"] = "West"
    if "base_price" not in df.columns:
        if "mrp" in df.columns:
            df["base_price"] = pd.to_numeric(df["mrp"], errors="coerce")
        elif "local_price" in df.columns:
            df["base_price"] = pd.to_numeric(df["local_price"], errors="coerce")
    if "margin" not in df.columns:
        if "gross_margin_pct" in df.columns:
            df["margin"] = pd.to_numeric(df["gross_margin_pct"], errors="coerce") * 100.0
        elif "profit_amt" in df.columns and "revenue" in df.columns:
            rev = pd.to_numeric(df["revenue"], errors="coerce").replace(0, np.nan)
            prof = pd.to_numeric(df["profit_amt"], errors="coerce")
            df["margin"] = (prof / rev * 100.0).fillna(0.0)
    if "festival_lift" not in df.columns and "festival_impact" in df.columns:
        df["festival_lift"] = pd.to_numeric(df["festival_impact"], errors="coerce")
    if "revenue" not in df.columns:
        df["revenue"] = pd.to_numeric(df.get("final_price", 0), errors="coerce").fillna(0) * pd.to_numeric(df.get("units_sold", 0), errors="coerce").fillna(0)
    if "cost" not in df.columns and "purchase_cost" in df.columns:
        df["cost"] = pd.to_numeric(df["purchase_cost"], errors="coerce")
    if "purchase_cost" not in df.columns and "local_price" in df.columns and "margin" in df.columns:
        lp = pd.to_numeric(df["local_price"], errors="coerce").fillna(0)
        mg = pd.to_numeric(df["margin"], errors="coerce").fillna(12.0)
        df["purchase_cost"] = lp * (1 - mg / 100.0)
    if "tax_pct" not in df.columns:
        df["tax_pct"] = 5.0
    if "mrp" not in df.columns and "local_price" in df.columns:
        df["mrp"] = pd.to_numeric(df["local_price"], errors="coerce").fillna(0) * 1.10
    if "lead_time_days" not in df.columns:
        df["lead_time_days"] = 3
    if "supplier_id" not in df.columns:
        df["supplier_id"] = "SUP000"
    if "net_profit" not in df.columns:
        df["net_profit"] = pd.to_numeric(df.get("profit_amt", 0), errors="coerce").fillna(0)
    if "is_festival_day" not in df.columns:
        df["is_festival_day"] = 0
    if "is_pre_festival_window" not in df.columns:
        df["is_pre_festival_window"] = 0
    if "hour_start" not in df.columns:
        df["hour_start"] = df["date"].dt.hour if np.issubdtype(df["date"].dtype, np.datetime64) else 12
    if "is_weekend" not in df.columns:
        df["is_weekend"] = (df["date"].dt.weekday >= 5).astype(int)
    numeric_cols = [
        "local_price","competitor_price","promo_depth_pct","final_price","units_sold","inventory_level","is_stockout",
        "revenue","margin","festival_lift","lead_time_days","purchase_cost","mrp","tax_pct","net_profit","cost",
        "opening_stock","stock_on_hand","reorder_point","needs_reorder","promo_flag","is_festival_day","is_pre_festival_window"
    ]
    for col in numeric_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce")
            if pd.api.types.is_float_dtype(df[col]):
                df[col] = df[col].astype("float32")
            elif pd.api.types.is_integer_dtype(df[col]):
                df[col] = df[col].astype("int32")
    for col in ["sku_id", "item_name", "category", "channel", "region", "festival_name", "payment_mode", "season", "weather_tag", "supplier_id"]:
        if col in df.columns:
            try:
                df[col] = df[col].astype("string")
            except Exception:
                pass
    console.print(f"Loaded {len(df)} rows from {path}", markup=False)
    req_cols = ["date","sku_id","category","channel","region","base_price","local_price","competitor_price","promo_flag","promo_depth_pct","final_price","units_sold","inventory_level","is_stockout","revenue","margin","festival_lift","seasonality_index","supplier_id","lead_time_days","purchase_cost","mrp","tax_pct","net_profit","cost"]
    for c in req_cols:
        if c not in df.columns:
            df[c] = np.nan
    try:
        if df["seasonality_index"].nunique(dropna=True) <= 1:
            console.print("Seasonality index constant or missing -> recomputing monthly seasonality index from units_sold (dataset-wide)...", markup=False)
            df["month"] = df["date"].dt.month
            monthly_avg = df.groupby("month")["units_sold"].mean().fillna(0.0)
            if monthly_avg.sum() == 0:
                default_idx = {m:1.0 for m in range(1,13)}
                df["seasonality_index"] = df["date"].dt.month.map(default_idx)
            else:
                monthly_idx = (monthly_avg / monthly_avg.mean()).to_dict()
                monthly_idx = {k: max(0.2, min(v, 5.0)) for k, v in monthly_idx.items()}
                df["seasonality_index"] = df["date"].dt.month.map(monthly_idx).fillna(1.0)
            df.drop(columns=["month"], inplace=True, errors=True)
            console.print("Seasonality enrichment applied and stored in memory (not overwriting CSV).", markup=False)
    except Exception as e:
        console.print(f"Seasonality enrichment failed: {e}", markup=False)

    # Time-series demand features for richer ML/DL behavior.
    try:
        time_order_cols = ["date"]
        if "txn_timestamp" in df.columns:
            df["txn_timestamp"] = pd.to_datetime(df["txn_timestamp"], errors="coerce")
            time_order_cols.append("txn_timestamp")
        df = df.sort_values(["sku_id"] + time_order_cols).reset_index(drop=True)
        daily = df.groupby(["sku_id", "date"], as_index=False).agg(
            daily_units=("units_sold", "sum"),
            avg_price=("local_price", "mean"),
            avg_festival=("festival_lift", "max"),
        )
        daily = daily.sort_values(["sku_id", "date"])
        daily["recent_7d_avg_units"] = daily.groupby("sku_id")["daily_units"].transform(lambda s: s.rolling(7, min_periods=1).mean())
        daily["recent_28d_avg_units"] = daily.groupby("sku_id")["daily_units"].transform(lambda s: s.rolling(28, min_periods=1).mean())
        daily["demand_trend_7_28"] = (daily["recent_7d_avg_units"] / (daily["recent_28d_avg_units"] + 1e-9) - 1.0).clip(-1.0, 2.0)
        df = df.merge(
            daily[["sku_id", "date", "recent_7d_avg_units", "recent_28d_avg_units", "demand_trend_7_28"]],
            on=["sku_id", "date"],
            how="left",
        )
    except Exception as e:
        console.print(f"Demand feature enrichment failed: {e}", markup=False)
    return df

def generate_synthetic_dataset(path: str, rows: int = 720) -> None:
    rng = np.random.default_rng(42)
    sku_meta = [
        ("GCPL_SKU_0001", "Soap", 192.11),
        ("GCPL_SKU_0002", "Soap", 196.03),
        ("GCPL_SKU_0003", "Insecticide", 190.08),
        ("GCPL_SKU_0004", "Hair Color", 185.26),
    ]
    channels = ["GT", "MT", "ECOM", "QCOMM"]
    regions = ["North", "South", "East", "West"]
    start_date = pd.Timestamp.now().normalize() - pd.Timedelta(days=rows // 4)

    out = []
    for i in range(rows):
        date = start_date + pd.Timedelta(days=i // 4)
        sku_id, category, base_price = sku_meta[i % len(sku_meta)]
        channel = channels[i % len(channels)]
        region = regions[(i // 3) % len(regions)]
        festival_lift = 1.0 + (0.18 if date.month in (3, 10, 11) else 0.0)
        seasonality_index = 1.0 + 0.08 * math.sin((date.month / 12) * 2 * math.pi)
        promo_flag = 1 if rng.random() < 0.28 else 0
        promo_depth = float(rng.uniform(4, 22)) if promo_flag else 0.0
        local_price = base_price * float(rng.uniform(0.96, 1.05))
        competitor_price = local_price * float(rng.uniform(0.95, 1.06))
        final_price = local_price * (1 - promo_depth / 100.0)
        demand_base = 95 + (12 if channel in ("ECOM", "QCOMM") else 0)
        units_sold = max(5.0, demand_base * festival_lift * seasonality_index * float(rng.uniform(0.75, 1.3)))
        inventory_level = max(8.0, units_sold * float(rng.uniform(10, 36)) / 30.0)
        is_stockout = 1 if inventory_level < (units_sold / 6.0) else 0
        revenue = final_price * units_sold
        margin_pct = float(rng.uniform(12, 32))
        out.append(
            {
                "date": date,
                "sku_id": sku_id,
                "category": category,
                "channel": channel,
                "region": region,
                "base_price": round(base_price, 2),
                "local_price": round(local_price, 2),
                "competitor_price": round(competitor_price, 2),
                "promo_flag": promo_flag,
                "promo_depth_pct": round(promo_depth, 2),
                "final_price": round(final_price, 2),
                "units_sold": round(units_sold, 2),
                "inventory_level": round(inventory_level, 2),
                "is_stockout": is_stockout,
                "revenue": round(revenue, 2),
                "margin": round(margin_pct, 2),
                "festival_lift": round(festival_lift, 3),
                "seasonality_index": round(seasonality_index, 3),
            }
        )

    pd.DataFrame(out).to_csv(path, index=False)
    console.print(f"Synthetic dataset created: {path} ({rows} rows)", markup=False)

# ---------------- Elasticity estimation (log-log OLS) ----------------
def estimate_elasticity(df: pd.DataFrame, sku=None, category=None) -> float:
    def loglog_slope(x, y):
        try:
            xm = np.log1p(np.array(x).astype(float))
            ym = np.log1p(np.array(y).astype(float))
            A = np.vstack([xm, np.ones_like(xm)]).T
            sol, *_ = np.linalg.lstsq(A, ym, rcond=None)
            slope = sol[0]
            return float(slope)
        except Exception:
            return None
    if sku is not None:
        sdf = df[df["sku_id"].astype(str) == str(sku)]
        if len(sdf) >= 10:
            e = loglog_slope(sdf["local_price"].replace(0,np.nan).dropna(), sdf["units_sold"].replace(0,np.nan).dropna())
            if e is not None and not math.isnan(e) and abs(e) > 0:
                return max(min(e, -0.1), -5.0)
    if category is not None:
        cdf = df[df["category"].astype(str) == str(category)]
        if len(cdf) >= 30:
            e = loglog_slope(cdf["local_price"].replace(0,np.nan).dropna(), cdf["units_sold"].replace(0,np.nan).dropna())
            if e is not None and not math.isnan(e) and abs(e) > 0:
                return max(min(e, -0.1), -5.0)
    gdf = df.dropna(subset=["local_price","units_sold"])
    if len(gdf) >= 50:
        e = loglog_slope(gdf["local_price"], gdf["units_sold"])
        if e is not None and not math.isnan(e) and abs(e) > 0:
            return max(min(e, -0.1), -5.0)
    return -0.8

# ---------------- AgentTrace dataclass ----------------
@dataclass
class AgentTrace:
    name: str
    candidates: Dict[str, float] = field(default_factory=dict)
    details: Dict[str, Any] = field(default_factory=dict)

# ---------------- Agents ----------------
class BasePriceAgent:
    def run(self, row, df_all) -> AgentTrace:
        prev_median = None
        sku = str(row.get("sku_id"))
        try:
            sku_hist = df_all[df_all["sku_id"].astype(str) == sku]
            if len(sku_hist) > 0:
                prev_median = float(sku_hist["local_price"].median(skipna=True))
        except Exception:
            prev_median = None
        if prev_median is None or math.isnan(prev_median):
            try:
                cat = str(row.get("category"))
                ch = str(row.get("channel"))
                rg = str(row.get("region"))
                grp = df_all[(df_all["category"].astype(str) == cat) & (df_all["channel"].astype(str) == ch) & (df_all["region"].astype(str) == rg)]
                if len(grp) > 0:
                    prev_median = float(grp["local_price"].median(skipna=True))
            except Exception:
                prev_median = None
        if prev_median is None or math.isnan(prev_median):
            prev_median = float(df_all["local_price"].median(skipna=True))
        seasonality = safe_float(row.get("seasonality_index"), 1.0)
        seasonality = max(0.2, min(seasonality, 5.0))
        festival = safe_float(row.get("festival_lift"), 1.0)
        neutral = prev_median * seasonality * festival
        conservative = neutral * 0.98
        aggressive = neutral * 1.05
        return AgentTrace("base", {"conservative": round(conservative,2), "neutral": round(neutral,2), "aggressive": round(aggressive,2)}, {"prev_median": prev_median, "seasonality": seasonality, "festival": festival})

class PromoAgent:
    CHANNEL_DEFAULT = {"ECOM": 20.0, "GT": 5.0, "MT": 12.0}
    def __init__(self, guardrails: Dict[str, Any]):
        self.guard = guardrails
    def compute_channel_default(self, channel: str) -> float:
        if not channel:
            return 10.0
        return float(self.CHANNEL_DEFAULT.get(channel.upper(), 10.0))
    def run(self, row, base_candidates: Dict[str, float]) -> AgentTrace:
        hist_promo = safe_float(row.get("promo_depth_pct"), None)
        channel = str(row.get("channel") or "").strip()
        comp_price = safe_float(row.get("competitor_price"), None)
        elasticity = safe_float(row.get("estimated_elasticity"), None)
        if elasticity is None or math.isnan(elasticity):
            elasticity = -0.8
        max_p = float(self.guard.get("max_promo_depth_pct", 40.0))
        base_pct = hist_promo if (hist_promo is not None and 0.0 <= hist_promo <= max_p) else self.compute_channel_default(channel)
        channel_mult = {"ECOM": 1.0, "GT": 0.5, "MT": 0.75}.get(channel.upper(), 0.9)
        e_factor = min(2.0, max(0.5, abs(elasticity) / 0.4))
        comp_gap_factor = 1.0
        if comp_price is not None and not math.isnan(comp_price):
            neutral_price = base_candidates.get("neutral", None)
            if neutral_price:
                gap = pct_change(neutral_price, comp_price)
                if gap < -5.0:
                    comp_gap_factor = 1.2 + min(0.8, abs(gap)/50.0)
                elif gap > 5.0:
                    comp_gap_factor = 0.9
        cand = {}
        for lvl, base_price in base_candidates.items():
            depth = base_pct * channel_mult * (1.0 + (e_factor - 1.0) * (1.0 if lvl == "aggressive" else 0.6))
            depth = depth * comp_gap_factor
            if lvl == "conservative":
                depth = depth * 0.5
            elif lvl == "neutral":
                depth = depth * 1.0
            else:
                depth = min(max_p, depth * 1.25)
            depth = max(0.0, min(depth, max_p))
            price_after = round(base_price * (1.0 - depth/100.0), 2)
            min_margin = self.guard.get("min_margin_pct", 12.0)
            cost = safe_float(row.get("cost"), None)
            if cost is None or math.isnan(cost):
                hist_price = safe_float(row.get("local_price"), base_price)
                hist_margin_pct = safe_float(row.get("margin"), None)
                if hist_margin_pct is None or math.isnan(hist_margin_pct):
                    cost = hist_price * (1.0 - min_margin/100.0)
                else:
                    cost = hist_price * (1.0 - hist_margin_pct/100.0)
            margin_pct_after = (price_after - cost) / (price_after + 1e-9) * 100.0 if price_after > 0 else -999.0
            if margin_pct_after < min_margin:
                required_price = cost / (1 - min_margin/100.0 + 1e-9)
                if required_price > price_after:
                    price_after = round(required_price, 2)
                    depth = max(0.0, (1.0 - price_after / (base_price + 1e-9)) * 100.0)
            cand[lvl] = round(price_after, 2)
        return AgentTrace("promo", cand, {"applied_base_pct": base_pct, "channel": channel, "elasticity": elasticity, "max_allowed": max_p})

class CompetitorAgent:
    def run(self, row, input_candidates):
        comp_price = safe_float(row.get("competitor_price"), None)
        out = {}
        if comp_price is None or math.isnan(comp_price):
            for lvl, p in input_candidates.items():
                out[lvl] = p
            return AgentTrace("competitor", out, {"note": "no competitor data"})
        for lvl, price_in in input_candidates.items():
            gap_pct = pct_change(price_in, comp_price)
            if lvl == "conservative":
                out[lvl] = price_in
            elif lvl == "neutral":
                if gap_pct < -5.0:
                    out[lvl] = round(price_in + 0.5 * (comp_price - price_in), 2)
                elif gap_pct > 5.0:
                    out[lvl] = round(min(price_in * 1.02, price_in + 2.0), 2)
                else:
                    out[lvl] = price_in
            else:
                if gap_pct < -5.0:
                    out[lvl] = round(price_in + 0.8 * (comp_price - price_in), 2)
                elif gap_pct > 5.0:
                    out[lvl] = round(min(price_in * 1.03, price_in + 3.0), 2)
                else:
                    out[lvl] = price_in
        return AgentTrace("competitor", out, {"comp_price": comp_price})

class InventoryAgent:
    def __init__(self, thresholds=None):
        self.thresholds = thresholds or {"low": 10, "high": 60}
    def infer_inventory_days(self, row) -> Optional[float]:
        if "inventory_days" in getattr(row, "index", []) and not pd.isna(row.get("inventory_days")):
            return safe_float(row.get("inventory_days"), None)
        inv_level = safe_float(row.get("inventory_level"), None)
        units_sold = safe_float(row.get("units_sold"), None)
        if inv_level is None:
            return None
        if units_sold is not None and units_sold > 0:
            daily = units_sold / 30.0
            if daily > 0:
                return inv_level / daily
        if inv_level > 200:
            return 90.0
        if inv_level > 100:
            return 60.0
        if inv_level > 30:
            return 30.0
        return 10.0
    def run(self, row, input_candidates):
        inv_days = self.infer_inventory_days(row)
        out = {}
        for lvl, price in input_candidates.items():
            if inv_days is None:
                out[lvl] = price
            else:
                if inv_days < self.thresholds["low"]:
                    if lvl == "conservative":
                        out[lvl] = round(price * 1.02, 2)
                    elif lvl == "neutral":
                        out[lvl] = round(price * 1.04, 2)
                    else:
                        out[lvl] = round(price * 1.07, 2)
                elif inv_days > self.thresholds["high"]:
                    if lvl == "conservative":
                        out[lvl] = round(price * 0.995, 2)
                    elif lvl == "neutral":
                        out[lvl] = round(price * 0.97, 2)
                    else:
                        out[lvl] = round(price * 0.94, 2)
                else:
                    out[lvl] = price
        return AgentTrace("inventory", out, {"inventory_days": inv_days})

class ProcurementAgent:
    def run(self, row, input_candidates):
        lead_time = safe_float(row.get("lead_time_days"), None)
        inv_level = safe_float(row.get("inventory_level"), None)
        units_sold = safe_float(row.get("units_sold"), None)
        if lead_time is None or math.isnan(lead_time):
            lead_time = 3.0
        inv_days = None
        if inv_level is not None and units_sold is not None and units_sold > 0:
            inv_days = inv_level / max(units_sold / 30.0, 1e-6)
        out = {}
        for lvl, price in input_candidates.items():
            adj = 1.0
            if lead_time >= 6 and (inv_days is None or inv_days < 18):
                adj = 1.015 if lvl == "conservative" else (1.025 if lvl == "neutral" else 1.04)
            elif lead_time <= 2 and (inv_days is not None and inv_days > 45):
                adj = 0.997 if lvl == "conservative" else (0.985 if lvl == "neutral" else 0.97)
            out[lvl] = round(price * adj, 2)
        return AgentTrace("procurement", out, {"lead_time_days": lead_time, "inventory_days": inv_days})

class BillingAgent:
    def run(self, row, input_candidates):
        mrp = safe_float(row.get("mrp"), None)
        purchase_cost = safe_float(row.get("purchase_cost"), None)
        tax_pct = safe_float(row.get("tax_pct"), 0.0)
        out = {}
        for lvl, price in input_candidates.items():
            p = float(price)
            if mrp is not None and not math.isnan(mrp) and mrp > 0:
                p = min(p, mrp)
            if purchase_cost is not None and not math.isnan(purchase_cost):
                floor = purchase_cost * (1.0 + max(0.01, min(tax_pct, 28.0) / 300.0))
                p = max(p, floor)
            out[lvl] = round(p, 2)
        return AgentTrace("billing", out, {"mrp": mrp, "tax_pct": tax_pct})

# ---------------- Candidate combine / mapping ----------------
class CandidateAssembler:
    def __init__(self, guardrails):
        self.guard = guardrails
    def assemble(self, traces):
        weights = {"base": 0.22, "promo": 0.24, "competitor": 0.20, "inventory": 0.14, "procurement": 0.10, "billing": 0.10}
        levels = ["conservative", "neutral", "aggressive"]
        combined = {lvl: 0.0 for lvl in levels}
        for lvl in levels:
            s = 0.0; tw = 0.0
            for aname, w in weights.items():
                trace = traces.get(aname)
                price = trace.candidates.get(lvl) if trace is not None else None
                if price is None: continue
                s += price * w; tw += w
            combined[lvl] = round(s / tw if tw > 0 else 0.0, 2)
        mapping = {"price_base": combined["conservative"], "price_optimal": combined["neutral"], "price_aggressive": combined["aggressive"]}
        return mapping

# ---------------- NN Demand Model (unchanged) ----------------
if ENABLE_NN:
    class DemandMLP(nn.Module):
        def __init__(self, n_in):
            super().__init__()
            self.net = nn.Sequential(
                nn.Linear(n_in, 128),
                nn.ReLU(),
                nn.Linear(128, 64),
                nn.ReLU(),
                nn.Linear(64,1)
            )
        def forward(self, x):
            return self.net(x)
    def train_nn_demand(csv_path=DATA_CSV, model_out="nn_demand_model.pt", scaler_out="nn_scaler.npz", epochs=8):
        from sklearn.preprocessing import StandardScaler
        df = load_and_enrich(csv_path)
        df = df.fillna(0)
        df['log_units'] = np.log1p(df['units_sold'].fillna(0))
        numeric = [
            'local_price','competitor_price','promo_depth_pct','inventory_level','festival_lift','seasonality_index',
            'lead_time_days','purchase_cost','mrp','tax_pct','recent_7d_avg_units','recent_28d_avg_units',
            'demand_trend_7_28','is_weekend','is_festival_day','is_pre_festival_window','hour_start'
        ]
        for c in numeric:
            if c not in df.columns:
                df[c] = 0.0
        df = pd.get_dummies(df, columns=['channel','region','category'], drop_first=True)
        feature_cols = [c for c in df.columns if c in numeric or c.startswith('channel_') or c.startswith('region_') or c.startswith('category_')]
        X = df[feature_cols].values.astype(float)
        y = df['log_units'].values.astype(float).reshape(-1,1)
        scaler = StandardScaler().fit(X)
        Xs = scaler.transform(X)
        split = int(0.9 * len(Xs))
        X_train, X_val = Xs[:split], Xs[split:]
        y_train, y_val = y[:split], y[split:]
        import torch.utils.data as data_utils
        train_ds = data_utils.TensorDataset(torch.tensor(X_train, dtype=torch.float32), torch.tensor(y_train, dtype=torch.float32))
        val_ds = data_utils.TensorDataset(torch.tensor(X_val, dtype=torch.float32), torch.tensor(y_val, dtype=torch.float32))
        train_loader = data_utils.DataLoader(train_ds, batch_size=2048, shuffle=True)
        val_loader = data_utils.DataLoader(val_ds, batch_size=2048, shuffle=False)
        model = DemandMLP(X_train.shape[1])
        opt = optim.Adam(model.parameters(), lr=1e-3)
        loss_fn = nn.MSELoss()
        best_val = 1e9
        for epoch in range(1, epochs+1):
            model.train()
            train_loss = 0.0
            for xb, yb in train_loader:
                pred = model(xb)
                loss = loss_fn(pred, yb)
                opt.zero_grad()
                loss.backward()
                opt.step()
                train_loss += loss.item() * xb.size(0)
            train_loss /= len(train_loader.dataset)
            model.eval()
            val_loss = 0.0
            with torch.no_grad():
                for xb, yb in val_loader:
                    val_loss += loss_fn(model(xb), yb).item() * xb.size(0)
            val_loss /= len(val_loader.dataset)
            console.print(f"Epoch {epoch}/{epochs} train_mse={train_loss:.6f} val_mse={val_loss:.6f}", markup=False)
            if val_loss < best_val:
                best_val = val_loss
                torch.save({"model_state": model.state_dict(), "feature_cols": feature_cols}, model_out)
                np.savez(scaler_out, mean=scaler.mean_, scale=scaler.scale_, feature_cols=np.array(feature_cols, dtype=object))
                console.print(f"Saved model to {model_out} and scaler to {scaler_out}", markup=False)
        return model_out, scaler_out
    def load_nn_model(model_path="nn_demand_model.pt", scaler_path="nn_scaler.npz"):
        if not os.path.exists(model_path) or not os.path.exists(scaler_path):
            raise FileNotFoundError("NN model or scaler not found. Train model first.")
        data = torch.load(model_path, map_location=torch.device("cpu"))
        feature_cols = data.get("feature_cols", None)
        model = DemandMLP(len(feature_cols))
        model.load_state_dict(data["model_state"])
        model.eval()
        scaler_np = np.load(scaler_path, allow_pickle=True)
        scaler = {"mean": scaler_np["mean"], "scale": scaler_np["scale"], "feature_cols": list(scaler_np["feature_cols"])}
        return model, scaler
    def predict_units_nn(row, model, scaler):
        rowd = row.to_dict()
        numeric = [
            'local_price','competitor_price','promo_depth_pct','inventory_level','festival_lift','seasonality_index',
            'lead_time_days','purchase_cost','mrp','tax_pct','recent_7d_avg_units','recent_28d_avg_units',
            'demand_trend_7_28','is_weekend','is_festival_day','is_pre_festival_window','hour_start'
        ]
        for c in numeric:
            if c not in rowd or rowd[c] is None:
                rowd[c] = 0.0
        feat = []
        for c in scaler["feature_cols"]:
            if c in numeric:
                feat.append(float(rowd.get(c, 0.0)))
            elif c.startswith("channel_") or c.startswith("region_") or c.startswith("category_"):
                base, _, suffix = c.partition("_")
                val = str(rowd.get(base,"")).strip().lower()
                feat.append(1.0 if val == suffix.strip().lower() else 0.0)
            else:
                feat.append(0.0)
        arr = np.array(feat).reshape(1,-1)
        arr_scaled = (arr - scaler["mean"]) / (scaler["scale"] + 1e-9)
        xt = torch.tensor(arr_scaled, dtype=torch.float32)
        with torch.no_grad():
            pred_log = model(xt).numpy().flatten()[0]
        pred_log = float(np.clip(pred_log, -20.0, 20.0))
        units = float(np.expm1(pred_log))
        units = max(0.0, min(units, 1e7))
        return units
else:
    def train_nn_demand(*args, **kwargs):
        raise RuntimeError("NN disabled")
    def load_nn_model(*args, **kwargs):
        raise RuntimeError("NN disabled")
    def predict_units_nn(*args, **kwargs):
        raise RuntimeError("NN disabled")

# ---------------- RL policy (unchanged) ----------------
if ENABLE_RL:
    def train_policy_simple(csv_path=DATA_CSV, policy_out="policy_rf.pkl", sample_size=30000):
        if not os.path.exists(csv_path):
            raise FileNotFoundError("CSV not found: " + csv_path)
        df = load_and_enrich(csv_path)
        df = df.sample(n=min(sample_size, len(df)), random_state=42).reset_index(drop=True)
        contexts = []
        labels = []
        for _, row in df.iterrows():
            b = BasePriceAgent().run(row, df)
            p = PromoAgent(DEFAULT_GUARDRAILS).run(row, b.candidates)
            c = CompetitorAgent().run(row, p.candidates)
            i = InventoryAgent().run(row, c.candidates)
            pr = ProcurementAgent().run(row, i.candidates)
            bl = BillingAgent().run(row, pr.candidates)
            mapped = CandidateAssembler(DEFAULT_GUARDRAILS).assemble({"base":b,"promo":p,"competitor":c,"inventory":i,"procurement":pr,"billing":bl})
            base_p = mapped["price_base"]
            opt_p = mapped["price_optimal"]
            agg_p = mapped["price_aggressive"]
            candidates = [base_p, opt_p, agg_p]
            rewards = []
            for cand in candidates:
                hist_units = safe_float(row.get("units_sold"), 0.0)
                hist_price = safe_float(row.get("local_price"), base_p)
                units = hist_units * (hist_price / (cand + 1e-9)) if hist_units>0 else 1.0
                purchase_cost = safe_float(row.get("purchase_cost"), None)
                margin_pct = safe_float(row.get("margin"), None)
                if purchase_cost is not None and not math.isnan(purchase_cost):
                    cost_est = purchase_cost
                else:
                    cost_est = cand * (1 - margin_pct/100.0) if (margin_pct and not math.isnan(margin_pct)) else cand * 0.8
                reward = (cand - cost_est) * max(0.0, units)
                rewards.append(reward)
            best_idx = int(np.argmax(rewards))
            ctx = [safe_float(row.get("local_price"),0.0), safe_float(row.get("competitor_price"),0.0), safe_float(row.get("promo_depth_pct"),0.0), safe_float(row.get("inventory_level"),0.0), safe_float(row.get("festival_lift"),0.0), safe_float(row.get("seasonality_index"),1.0), safe_float(row.get("lead_time_days"),3.0), safe_float(row.get("purchase_cost"),0.0), safe_float(row.get("mrp"),0.0), safe_float(row.get("tax_pct"),5.0)]
            contexts.append(ctx)
            labels.append(best_idx)
        X = np.array(contexts); y = np.array(labels)
        clf = RandomForestClassifier(n_estimators=100, random_state=42, n_jobs=-1)
        clf.fit(X, y)
        joblib.dump(clf, policy_out)
        console.print(f"Saved policy to {policy_out}", markup=False)
        return policy_out
    def load_policy(policy_path="policy_rf.pkl"):
        if not os.path.exists(policy_path):
            raise FileNotFoundError("Policy file not found: " + policy_path)
        clf = joblib.load(policy_path)
        return clf
else:
    def train_policy_simple(*args, **kwargs):
        raise RuntimeError("RL disabled")
    def load_policy(*args, **kwargs):
        raise RuntimeError("RL disabled")

# ---------------- Predictor for units ----------------
def predict_units_for_candidate(row, price, elasticity, nn_model=None, nn_scaler=None, demand_prior=None):
    hist_units = safe_float(row.get("units_sold"), 0.0)
    hist_price = safe_float(row.get("local_price"), price)
    recent_7 = safe_float(row.get("recent_7d_avg_units"), float("nan"))
    recent_28 = safe_float(row.get("recent_28d_avg_units"), float("nan"))
    trend_7_28 = safe_float(row.get("demand_trend_7_28"), 0.0)
    if hist_units <= 0 and not math.isnan(recent_7):
        hist_units = max(hist_units, recent_7)

    units_elasticity = None
    if hist_units > 0:
        units_elasticity = hist_units * ((hist_price + 1e-9)/(price + 1e-9)) ** max(abs(elasticity), 0.1)
        units_elasticity = float(max(0.0, min(units_elasticity, 1e7)))

    units_nn = None
    if ENABLE_NN and (nn_model is not None and nn_scaler is not None):
        try:
            units = predict_units_nn(row, nn_model, nn_scaler)
            if hist_price > 0:
                units = units * (hist_price / (price + 1e-9)) ** max(abs(elasticity), 0.1)
            units_nn = float(max(0.0, min(units, 1e7)))
        except Exception:
            units_nn = None

    if units_elasticity is None:
        if hist_units <= 0:
            hist_units = max(1.0, safe_float(row.get("revenue"),0.0)/(hist_price+1e-9))
        units_elasticity = hist_units * ((hist_price + 1e-9)/(price + 1e-9)) ** max(abs(elasticity), 0.1)
        units_elasticity = float(max(0.0, min(units_elasticity, 1e7)))

    # DeepAR-style prior blended in if provided by forecast path.
    units_temporal = None
    if demand_prior is not None:
        try:
            units_temporal = float(max(0.0, min(float(demand_prior), 1e7)))
            if hist_price > 0:
                units_temporal = units_temporal * ((hist_price + 1e-9)/(price + 1e-9)) ** max(abs(elasticity), 0.1)
        except Exception:
            units_temporal = None

    units_recent = None
    if not math.isnan(recent_7):
        units_recent = float(recent_7 * (1.0 + max(-0.5, min(1.5, trend_7_28))))
    elif not math.isnan(recent_28):
        units_recent = float(recent_28)

    candidates = []
    if units_elasticity is not None:
        candidates.append(("elasticity", units_elasticity))
    if units_nn is not None:
        candidates.append(("nn", units_nn))
    if units_temporal is not None:
        candidates.append(("deepar_proxy", units_temporal))
    if units_recent is not None:
        candidates.append(("recent", units_recent))
    if not candidates:
        candidates.append(("fallback", max(1.0, safe_float(row.get("revenue"),0.0)/(price+1e-9))))

    weights = {"elasticity": 0.35, "nn": 0.35, "deepar_proxy": 0.20, "recent": 0.10, "fallback": 1.0}
    num = 0.0
    den = 0.0
    for name, val in candidates:
        w = weights.get(name, 0.1)
        num += w * val
        den += w
    units_new = num / den if den > 0 else candidates[0][1]
    return float(max(0.0, min(units_new, 1e7)))

# ---------------- LLM (Bedrock) helpers (safety enforced) ----------------
def _flatten_llm_messages(messages: list) -> Dict[str, str]:
    system_parts = []
    user_parts = []
    for msg in messages:
        role = str(msg.get("role", "user")).strip().lower()
        content = msg.get("content", "")
        if isinstance(content, list):
            joined = "\n".join(str(x) for x in content)
        else:
            joined = str(content)
        if role == "system":
            system_parts.append(joined)
        else:
            user_parts.append(f"{role.upper()}:\n{joined}")
    return {
        "system": "\n\n".join(p for p in system_parts if p.strip()),
        "user": "\n\n".join(p for p in user_parts if p.strip()),
    }


def _bedrock_request_for_model(model_id: str, messages: list, max_tokens: int, temperature: float) -> Dict[str, Any]:
    flat = _flatten_llm_messages(messages)
    system_text = flat["system"]
    user_text = flat["user"] or "Provide a concise response."

    if model_id.startswith("amazon.nova"):
        req = {
            "messages": [{"role": "user", "content": [{"text": user_text}]}],
            "inferenceConfig": {"max_new_tokens": max_tokens, "temperature": temperature, "topP": 0.9},
        }
        if system_text:
            req["system"] = [{"text": system_text}]
        return req

    if model_id.startswith("anthropic."):
        prompt_text = f"{system_text}\n\n{user_text}".strip() if system_text else user_text
        return {
            "anthropic_version": "bedrock-2023-05-31",
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": [{"role": "user", "content": prompt_text}],
        }

    return {
        "model": model_id,
        "messages": messages,
        "max_completion_tokens": max_tokens,
        "temperature": temperature,
        "top_p": 0.9,
        "stream": False,
    }


def _bedrock_parse_response(model_id: str, parsed: Dict[str, Any]) -> str:
    if model_id.startswith("amazon.nova"):
        outputs = parsed.get("output", {}).get("message", {}).get("content", [])
        texts = [str(x.get("text", "")) for x in outputs if isinstance(x, dict)]
        return "\n".join(t for t in texts if t).strip()

    if model_id.startswith("anthropic."):
        blocks = parsed.get("content", [])
        texts = [str(x.get("text", "")) for x in blocks if isinstance(x, dict)]
        return "\n".join(t for t in texts if t).strip()

    out_text = ""
    for choice in parsed.get("choices", []):
        msg = choice.get("message", {}).get("content")
        if isinstance(msg, list):
            out_text += "\n".join(str(x) for x in msg)
        elif msg:
            out_text += str(msg)
    return out_text or parsed.get("output", parsed.get("text", str(parsed)))


def bedrock_invoke(messages: list, model_id: str = BEDROCK_MODEL_ID, max_tokens: int = LLM_MAX_TOKENS, temperature: float = 0.0):
    if boto3 is None:
        raise RuntimeError("boto3 not installed/configured to call Bedrock.")
    client = boto3.client("bedrock-runtime", region_name=AWS_REGION)
    import json as _json

    model_candidates = [model_id] + [m for m in BEDROCK_FALLBACK_MODELS if m != model_id]
    last_error = None
    for candidate in model_candidates:
        try:
            native_request = _bedrock_request_for_model(candidate, messages, max_tokens=max_tokens, temperature=temperature)
            response = client.invoke_model(modelId=candidate, body=_json.dumps(native_request))
            body = response["body"].read().decode("utf-8")
            parsed = _json.loads(body)
            text = _bedrock_parse_response(candidate, parsed)
            if text:
                return text
            last_error = RuntimeError(f"Empty response from model {candidate}")
        except Exception as e:
            last_error = e
            continue
    raise RuntimeError(f"Bedrock invocation failed across models: {last_error}")


def _bedrock_ping_via_invoke(region: str) -> bool:
    """Minimal invoke to verify Bedrock works (e.g. when list_foundation_models is not allowed). Uses 1 token."""
    import json as _json
    runtime = boto3.client("bedrock-runtime", region_name=region)
    candidates = [BEDROCK_MODEL_ID] + [m for m in BEDROCK_FALLBACK_MODELS if m != BEDROCK_MODEL_ID]
    ping_messages = [{"role": "user", "content": "Say OK"}]
    for model_id in candidates:
        try:
            req = _bedrock_request_for_model(model_id, ping_messages, max_tokens=1, temperature=0.0)
            resp = runtime.invoke_model(modelId=model_id, body=_json.dumps(req))
            body = resp["body"].read().decode("utf-8")
            parsed = _json.loads(body)
            _bedrock_parse_response(model_id, parsed)
            return True
        except Exception:
            continue
    return False


def get_bedrock_status(force: bool = False) -> Dict[str, Any]:
    now = time.time()
    if not force and now - float(_BEDROCK_STATUS_CACHE.get("checked_at", 0.0)) < 60:
        return dict(_BEDROCK_STATUS_CACHE)
    status = {"checked_at": now, "ok": False, "error": ""}
    try:
        if boto3 is None:
            raise RuntimeError("boto3 is not installed on the backend.")
        # Use env region so terminal exports (AWS_DEFAULT_REGION / AWS_REGION) are respected
        region = os.getenv("AWS_REGION") or os.getenv("AWS_DEFAULT_REGION") or "ap-south-1"
        session = boto3.Session(region_name=region)
        creds = session.get_credentials()
        if creds is None:
            raise RuntimeError(
                "AWS credentials not found. In the same terminal where you start the backend, run: "
                "export AWS_ACCESS_KEY_ID=... ; export AWS_SECRET_ACCESS_KEY=... ; export AWS_DEFAULT_REGION=ap-south-1"
            )
        bedrock_client = session.client("bedrock", region_name=region)
        # Prefer list_foundation_models (needs bedrock:ListFoundationModels)
        try:
            resp = bedrock_client.list_foundation_models()
            models = resp.get("modelSummaries", []) if isinstance(resp, dict) else []
            if models:
                list_ok = True
                status["ok"] = True
                status["error"] = ""
                _BEDROCK_STATUS_CACHE.update(status)
                return dict(_BEDROCK_STATUS_CACHE)
        except Exception:
            pass
        # Fallback: if list fails or returns empty, check via actual invoke (same path as price review)
        if _bedrock_ping_via_invoke(region):
            status["ok"] = True
            status["error"] = ""
        else:
            status["ok"] = False
            status["error"] = "List and invoke check failed. Check IAM: bedrock:ListFoundationModels and/or bedrock:InvokeModel."
    except Exception as exc:
        status["ok"] = False
        err = str(exc)
        status["error"] = err
        if "AccessDenied" in err or "Unauthorized" in err:
            status["error"] = f"{err} — Check IAM has bedrock:ListFoundationModels and bedrock:InvokeModel."
    _BEDROCK_STATUS_CACHE.update(status)
    return dict(_BEDROCK_STATUS_CACHE)

def build_explain_messages(context_meta: Dict[str,Any], subject: Dict[str,Any], agent_traces: Dict[str,Any], sop: Dict[str,Any]) -> list:
    system = {"role":"system", "content": "You are a concise pricing analyst. Provide a short clear explanation (max ~600 tokens). Do NOT propose prices that differ from the engine output. Only explain the engine output."}
    user = {"role":"user", "content": json.dumps({"context_meta":context_meta, "subject":subject, "agent_traces":agent_traces, "sop_guardrails":sop.get("guardrails", DEFAULT_GUARDRAILS)}, indent=2)}
    return [system, user]

# ---------------- Forecast function (unchanged improvement) ----------------
def _estimate_monthly_trend(df_all: pd.DataFrame, sku) -> float:
    try:
        sdf = df_all[df_all["sku_id"].astype(str) == str(sku)]
        if sdf.shape[0] < 6:
            return 0.0
        sdf = sdf.copy()
        sdf['month_index'] = (sdf['date'].dt.year - sdf['date'].dt.year.min())*12 + sdf['date'].dt.month
        monthly = sdf.groupby('month_index')['units_sold'].median().dropna()
        if len(monthly) < 3:
            return 0.0
        x = np.arange(len(monthly)).astype(float)
        y = monthly.values.astype(float)
        A = np.vstack([x, np.ones_like(x)]).T
        slope, intercept = np.linalg.lstsq(A, y, rcond=None)[0]
        mean_y = y.mean() if y.mean()!=0 else 1.0
        trend_pct = float(slope / (mean_y + 1e-9))
        trend_pct = max(-0.5, min(trend_pct, 0.5))
        return trend_pct
    except Exception:
        return 0.0

def _build_daily_sku_series(df_all: pd.DataFrame, sku: str) -> pd.DataFrame:
    sdf = df_all[df_all["sku_id"].astype(str) == str(sku)].copy()
    if sdf.empty:
        sdf = df_all.copy()
    daily = sdf.groupby("date", as_index=False).agg(
        daily_units=("units_sold", "sum"),
        local_price=("local_price", "mean"),
        promo_depth_pct=("promo_depth_pct", "mean"),
        festival_lift=("festival_lift", "max"),
        seasonality_index=("seasonality_index", "mean"),
    )
    daily = daily.sort_values("date")
    daily["dow"] = daily["date"].dt.weekday
    return daily

def deepar_sagemaker_forecast(
    df_all: pd.DataFrame,
    sku: str,
    start: pd.Timestamp,
    days: int,
    num_samples: int = 200,
    dataset_key: Optional[str] = None,
) -> Optional[List[Dict[str, float]]]:
    """
    Prefer SageMaker DeepAR endpoint when configured.
    Expected endpoint accepts DeepAR JSON payload.
    """
    endpoint_name = get_deepar_endpoint(dataset_key)
    if not endpoint_name or boto3 is None:
        return None
    try:
        daily = _build_daily_sku_series(df_all, sku)
        if daily.empty:
            return None
        target = daily["daily_units"].astype(float).tail(365).tolist()
        if len(target) < 14:
            return None
        sku_vocab = sorted(df_all["sku_id"].astype(str).unique().tolist())
        sku_cat = int(sku_vocab.index(str(sku))) if str(sku) in sku_vocab else 0
        hist_promo = daily["promo_depth_pct"].fillna(0.0).tail(365).astype(float).tolist() if "promo_depth_pct" in daily.columns else [0.0] * len(target)
        hist_festival = daily["festival_lift"].fillna(1.0).tail(365).astype(float).tolist() if "festival_lift" in daily.columns else [1.0] * len(target)
        hist_price = daily["local_price"].ffill().bfill().fillna(0.0).tail(365).astype(float).tolist() if "local_price" in daily.columns else [0.0] * len(target)
        calendar_specs = load_calendar_event_specs()
        future_dates = pd.date_range(start, periods=days)
        future_promo = []
        future_festival = []
        last_price = hist_price[-1] if hist_price else 0.0
        future_price = [float(last_price)] * days
        for fd in future_dates:
            fctx = festival_context_for_date(fd, calendar_specs)
            future_promo.append(float(fctx.get("promo_depth_pct") or 0.0))
            future_festival.append(float(fctx.get("festival_lift", 1.0)))
        start_hist = str(pd.to_datetime(daily["date"].min()).date()) + " 00:00:00"
        payload = {
            "instances": [{
                "start": start_hist,
                "target": target,
                "cat": [sku_cat],
                "dynamic_feat": [
                    hist_promo + future_promo,
                    hist_festival + future_festival,
                    hist_price + future_price,
                ],
            }],
            "configuration": {
                "num_samples": int(max(50, min(num_samples, 1000))),
                "output_types": ["mean", "quantiles"],
                "quantiles": ["0.1", "0.5", "0.9"],
                "prediction_length": int(days),
                "freq": "D",
            },
        }
        client = boto3.client("sagemaker-runtime", region_name=AWS_REGION)
        resp = client.invoke_endpoint(
            EndpointName=endpoint_name,
            ContentType="application/json",
            Body=json.dumps(payload).encode("utf-8"),
        )
        body = resp["Body"].read().decode("utf-8")
        parsed = json.loads(body)
        pred = parsed.get("predictions", [])
        if not pred:
            return None
        p0 = pred[0]
        means = p0.get("mean", []) or []
        q = p0.get("quantiles", {}) or {}
        q10 = q.get("0.1", [0.0] * len(means))
        q50 = q.get("0.5", means if means else [0.0] * len(q10))
        q90 = q.get("0.9", [0.0] * len(q50))
        n = min(days, len(q50))
        out = []
        for i in range(n):
            out.append({
                "mean": float(means[i]) if i < len(means) else float(q50[i]),
                "p10": float(q10[i]) if i < len(q10) else float(q50[i]),
                "p50": float(q50[i]),
                "p90": float(q90[i]) if i < len(q90) else float(q50[i]),
            })
        if len(out) < days and out:
            out.extend([out[-1]] * (days - len(out)))
        return out
    except Exception as e:
        console.print(f"DeepAR endpoint forecast failed: {e}. Falling back to local proxy.", markup=False)
        return None

_CALENDAR_STATE_CACHE: Dict[str, Dict[str, Any]] = {}


def load_calendar_event_specs() -> List[Dict[str, Any]]:
    try:
        parsed = json.loads(CALENDAR_EVENTS_JSON)
        return parsed if isinstance(parsed, list) else []
    except Exception as exc:
        console.print(f"Calendar event JSON parse failed: {exc}. Festival boosts disabled.", markup=False)
        return []


def _ssm_calendar_state(calendar_arn: str, at_time: pd.Timestamp) -> str:
    cache_key = f"{calendar_arn}|{str(at_time.normalize().date())}"
    cached = _CALENDAR_STATE_CACHE.get(cache_key)
    if cached:
        return str(cached.get("state", "OPEN"))
    if boto3 is None:
        return "OPEN"
    client = boto3.client("ssm", region_name=AWS_REGION)
    resp = client.get_calendar_state(CalendarNames=[calendar_arn], AtTime=f"{str(at_time.normalize().date())}T00:00:00Z")
    state = str(resp.get("State", "OPEN")).upper()
    _CALENDAR_STATE_CACHE[cache_key] = {"state": state}
    return state


def festival_context_for_date(d: pd.Timestamp, calendar_specs: Optional[List[Dict[str, Any]]] = None) -> Dict[str, Any]:
    specs = calendar_specs if calendar_specs is not None else load_calendar_event_specs()
    out = {
        "festival_name": "",
        "festival_lift": 1.0,
        "is_festival_day": 0,
        "is_pre_festival_window": 0,
        "promo_depth_pct": None,
    }
    if not specs:
        return out

    for ev in specs:
        calendar_arn = str(ev.get("calendar_arn", "")).strip()
        if not calendar_arn:
            continue
        name = str(ev.get("name", "Festival"))
        boost = float(ev.get("boost", 1.0))
        promo_depth = float(ev.get("promo_depth_pct", 0.0))
        pre_days = int(ev.get("pre_window_days", 4))
        pre_boost = float(ev.get("pre_window_boost", max(1.0, boost * 0.7)))
        try:
            state_now = _ssm_calendar_state(calendar_arn, d)
            if state_now == "CLOSED":
                out["festival_name"] = name
                out["festival_lift"] = max(out["festival_lift"], boost)
                out["is_festival_day"] = 1
                out["promo_depth_pct"] = max(float(out["promo_depth_pct"] or 0.0), promo_depth)
                return out
            for offset in range(1, max(pre_days, 0) + 1):
                future_date = d + pd.Timedelta(days=offset)
                state_future = _ssm_calendar_state(calendar_arn, future_date)
                if state_future == "CLOSED":
                    out["festival_name"] = f"Pre-{name}"
                    out["festival_lift"] = max(out["festival_lift"], pre_boost)
                    out["is_pre_festival_window"] = 1
                    out["promo_depth_pct"] = max(float(out["promo_depth_pct"] or 0.0), max(4.0, promo_depth * 0.65))
                    break
        except Exception as exc:
            console.print(f"Calendar lookup failed for {name}: {exc}", markup=False)
    return out

def deepar_proxy_forecast(df_all: pd.DataFrame, sku: str, start: pd.Timestamp, days: int, n_samples: int = 200) -> List[Dict[str, float]]:
    """
    Lightweight probabilistic autoregressive forecaster inspired by DeepAR behavior:
    global-ish seasonality, per-SKU autoregressive state, and Monte Carlo quantiles.
    """
    daily = _build_daily_sku_series(df_all, sku)
    if daily.empty:
        return [{"mean": 1.0, "p10": 1.0, "p50": 1.0, "p90": 1.0} for _ in range(days)]

    hist = daily["daily_units"].astype(float).values
    if len(hist) < 7:
        base = float(max(1.0, np.mean(hist) if len(hist) else 1.0))
        return [{"mean": base, "p10": 0.8 * base, "p50": base, "p90": 1.2 * base} for _ in range(days)]

    level = float(pd.Series(hist).ewm(alpha=0.25, adjust=False).mean().iloc[-1])
    resid = hist[1:] - pd.Series(hist).rolling(7, min_periods=1).mean().values[1:]
    sigma = float(np.nanstd(resid))
    sigma = max(0.05 * max(level, 1.0), sigma)

    daily["dow"] = daily["date"].dt.weekday
    dow_mean = daily.groupby("dow")["daily_units"].mean()
    global_mean = float(max(1e-6, daily["daily_units"].mean()))
    dow_factor = {int(k): float(v / global_mean) for k, v in dow_mean.items()}

    results = []
    rng = np.random.default_rng(abs(hash(str(sku))) % (2**32))
    simulated_hist = list(hist[-28:])
    for i, d in enumerate(pd.date_range(start, periods=days)):
        dow = int(d.weekday())
        sf = float(max(0.25, min(5.0, dow_factor.get(dow, 1.0))))
        festival_boost = float(festival_context_for_date(d).get("festival_lift", 1.0))

        sims = []
        base_recent = np.mean(simulated_hist[-7:]) if len(simulated_hist) >= 7 else np.mean(simulated_hist)
        for _ in range(n_samples):
            eps = float(rng.normal(0.0, sigma))
            ar = 0.6 * base_recent + 0.4 * level
            y = max(0.0, (ar * sf * festival_boost) + eps)
            sims.append(y)
        sims_arr = np.array(sims, dtype=float)
        p10 = float(np.percentile(sims_arr, 10))
        p50 = float(np.percentile(sims_arr, 50))
        p90 = float(np.percentile(sims_arr, 90))
        mean_v = float(sims_arr.mean())
        results.append({"mean": mean_v, "p10": p10, "p50": p50, "p90": p90})
        simulated_hist.append(mean_v)
        if len(simulated_hist) > 56:
            simulated_hist = simulated_hist[-56:]

    return results

def forecast_and_price(
    df_all,
    sop,
    sku,
    channel,
    region,
    start_date_str,
    days,
    nn_model=None,
    nn_scaler=None,
    policy=None,
    festival_context: Optional[Dict[str, Any]] = None,
    dataset_key: Optional[str] = None,
):
    sku = resolve_input_value(str(sku), sorted(df_all["sku_id"].astype(str).unique().tolist()), "SKU_ID")
    channel = resolve_input_value(str(channel), sorted(df_all["channel"].astype(str).unique().tolist()), "Channel")
    region = resolve_input_value(str(region), sorted(df_all["region"].astype(str).unique().tolist()), "Region")
    start = pd.to_datetime(start_date_str)
    sku_mask = df_all["sku_id"].astype(str) == str(sku)
    if sku_mask.sum() >= 5:
        hist = df_all[sku_mask].copy()
    else:
        hist = df_all.copy()
    latest_sku_row = hist.sort_values("date").iloc[-1].copy() if not hist.empty else _build_synthetic_row(df_all, sku=sku, channel=channel, region=region)
    monthly = df_all.groupby(df_all["date"].dt.month)["units_sold"].mean().fillna(0.0)
    if monthly.sum() == 0:
        monthly = pd.Series({m:1.0 for m in range(1,13)})
    monthly = monthly / (monthly.mean() + 1e-9)
    monthly = monthly.apply(lambda x: float(max(0.2, min(x, 5.0))))
    baseline_daily = max(1.0, hist["units_sold"].mean() / 30.0 if hist["units_sold"].mean()>0 else 1.0)
    sku_trend_monthly_pct = _estimate_monthly_trend(df_all, sku)
    results = []
    seed = abs(hash(str(sku))) % (10**8)
    rnd = random.Random(seed)
    deepar_endpoint = get_deepar_endpoint(dataset_key)
    prob_path = deepar_sagemaker_forecast(df_all, sku, start, days, num_samples=260, dataset_key=dataset_key)
    deepar_source = "sagemaker_endpoint" if prob_path is not None else "local_proxy"
    if prob_path is None:
        prob_path = deepar_proxy_forecast(df_all, sku, start, days, n_samples=220)
    calendar_specs = load_calendar_event_specs()
    for day_idx, d in enumerate(pd.date_range(start, periods=days)):
        m = int(d.month)
        season_multiplier = monthly.get(m, monthly.mean())
        months_ahead = ((d.year - pd.to_datetime(df_all["date"].max()).year) * 12 + (d.month - pd.to_datetime(df_all["date"].max()).month))
        months_ahead = max(0, months_ahead)
        trend_multiplier = 1.0 + sku_trend_monthly_pct * min(months_ahead, 6)
        noise = (rnd.uniform(-0.05, 0.05)) * (1 + months_ahead/12.0)
        model_pred = float(max(0.0, baseline_daily * season_multiplier * trend_multiplier * (1.0 + noise)))
        ar_pred = prob_path[day_idx]["mean"] if day_idx < len(prob_path) else model_pred
        fctx = festival_context_for_date(d, calendar_specs)
        festival_boost = float(fctx.get("festival_lift", 1.0))
        if festival_context and isinstance(festival_context, dict):
            festival_boost = max(festival_boost, float(festival_context.get("festival_multiplier", 1.0)))
        # DeepAR path is primary when available; trend model acts as stabilizer.
        w_ar = 0.80 if deepar_source == "sagemaker_endpoint" else 0.55
        pred_units = float((w_ar * ar_pred + (1.0 - w_ar) * model_pred) * festival_boost)
        base_category = (
            latest_sku_row.get("category")
            if str(latest_sku_row.get("sku_id", "")) == str(sku)
            else (df_all[df_all["sku_id"].astype(str) == str(sku)]["category"].mode().iloc[0] if (df_all[df_all["sku_id"].astype(str) == str(sku)].shape[0] > 0) else df_all["category"].mode().iloc[0])
        )
        current_inventory = safe_float(latest_sku_row.get("inventory_level"), safe_float(latest_sku_row.get("stock_on_hand"), 0.0))
        lead_time = safe_float(latest_sku_row.get("lead_time_days"), float(df_all["lead_time_days"].median(skipna=True)))
        recent7 = float(hist["units_sold"].tail(7).mean()) if len(hist) else pred_units
        recent28 = float(hist["units_sold"].tail(28).mean()) if len(hist) else pred_units
        reorder_point = max(recent7 * max(lead_time, 1.0) * 1.15, recent28 * 0.45)
        synthetic = {
            "sku_id": sku,
            "item_name": str(latest_sku_row.get("item_name", "")),
            "channel": channel or str(latest_sku_row.get("channel", df_all["channel"].mode().iloc[0])),
            "region": region or str(latest_sku_row.get("region", df_all["region"].mode().iloc[0])),
            "category": base_category,
            "supplier_id": str(latest_sku_row.get("supplier_id", "SUP000")),
            "local_price": safe_float(latest_sku_row.get("local_price"), float(df_all["local_price"].median(skipna=True))),
            "competitor_price": safe_float(latest_sku_row.get("competitor_price"), float(df_all["competitor_price"].median(skipna=True))),
            "promo_depth_pct": (
                float(festival_context.get("promo_depth_pct"))
                if (festival_context and festival_context.get("promo_depth_pct") is not None)
                else (float(fctx.get("promo_depth_pct")) if fctx.get("promo_depth_pct") is not None else safe_float(latest_sku_row.get("promo_depth_pct"), float(df_all["promo_depth_pct"].median(skipna=True))))
            ),
            "opening_stock": current_inventory,
            "stock_on_hand": current_inventory,
            "inventory_level": current_inventory,
            "reorder_point": reorder_point,
            "units_sold": float(pred_units * 30.0),
            "seasonality_index": float(season_multiplier),
            "festival_lift": festival_boost,
            "lead_time_days": lead_time,
            "purchase_cost": safe_float(latest_sku_row.get("purchase_cost"), float(df_all["purchase_cost"].median(skipna=True))),
            "mrp": safe_float(latest_sku_row.get("mrp"), float(df_all["mrp"].median(skipna=True))),
            "tax_pct": safe_float(latest_sku_row.get("tax_pct"), float(df_all["tax_pct"].median(skipna=True))),
            "recent_7d_avg_units": recent7,
            "recent_28d_avg_units": recent28,
            "demand_trend_7_28": float((recent7 / (recent28 + 1e-9) - 1.0)) if recent28 > 0 else 0.0,
            "is_weekend": int(d.weekday() >= 5),
            "is_festival_day": int(fctx.get("is_festival_day", 0)),
            "is_pre_festival_window": int(fctx.get("is_pre_festival_window", 0)),
            "hour_start": 12,
            "festival_name": (
                ",".join(festival_context.get("active_festivals", []))
                if (festival_context and isinstance(festival_context.get("active_festivals", None), list))
                else str(fctx.get("festival_name", ""))
            ),
        }
        row = pd.Series(synthetic)
        base_tr = BasePriceAgent().run(row, df_all)
        promo_tr = PromoAgent(sop.get("guardrails", DEFAULT_GUARDRAILS)).run(row, base_tr.candidates)
        comp_tr = CompetitorAgent().run(row, promo_tr.candidates)
        inv_tr = InventoryAgent().run(row, comp_tr.candidates)
        proc_tr = ProcurementAgent().run(row, inv_tr.candidates)
        bill_tr = BillingAgent().run(row, proc_tr.candidates)
        traces = {"base": base_tr, "promo": promo_tr, "competitor": comp_tr, "inventory": inv_tr, "procurement": proc_tr, "billing": bill_tr}
        mapped = CandidateAssembler(sop.get("guardrails", DEFAULT_GUARDRAILS)).assemble(traces)
        elasticity = estimate_elasticity(df_all, sku=sku, category=synthetic["category"])
        candidates_detail = {}
        for k, v in mapped.items():
            pred_units_k = predict_units_for_candidate(row, v, elasticity, nn_model, nn_scaler, demand_prior=ar_pred)
            pred_units_k = float(max(0.0, min(pred_units_k, 1e7)))
            candidates_detail[k] = {"price": round(v,2), "pred_units": pred_units_k, "elasticity": elasticity}
        selection = enrich_selection_payload(row, selector_and_sop(row, mapped, sop.get("guardrails", DEFAULT_GUARDRAILS)), candidates_detail)
        for k,v in candidates_detail.items():
            v["pred_units"] = safe_num_for_json(v.get("pred_units"), 0.0)
            v["price"] = safe_num_for_json(v.get("price"), None)
            v["elasticity"] = safe_num_for_json(v.get("elasticity"), None)
        results.append({
            "date": str(d.date()),
            "seasonality_index": float(season_multiplier),
            "sku_trend_monthly_pct": float(sku_trend_monthly_pct),
            "demand_quantiles": {
                "p10": float(prob_path[day_idx]["p10"]) if day_idx < len(prob_path) else pred_units,
                "p50": float(prob_path[day_idx]["p50"]) if day_idx < len(prob_path) else pred_units,
                "p90": float(prob_path[day_idx]["p90"]) if day_idx < len(prob_path) else pred_units,
                "mean": float(prob_path[day_idx]["mean"]) if day_idx < len(prob_path) else pred_units,
            },
            "forecast_engine": deepar_source,
            "deepar_endpoint": deepar_endpoint if deepar_source == "sagemaker_endpoint" else "",
            "traces": {k: {lvl: safe_num_for_json(p, None) for lvl,p in t.candidates.items()} for k,t in traces.items()},
            "candidates": candidates_detail,
            "selection": {k: safe_num_for_json(v, None) for k,v in selection.items()}
        })
    return results

# ---------------- SOP selection & enforcement (unchanged) ----------------
def selector_and_sop(row, candidate_prices, guardrails):
    min_margin = guardrails.get("min_margin_pct", 12.0)
    cost = safe_float(row.get("cost"), None)
    if cost is None or math.isnan(cost):
        cost = safe_float(row.get("purchase_cost"), None)
    if cost is None or math.isnan(cost):
        prev = safe_float(row.get("local_price"), candidate_prices.get("price_base", 100.0))
        cost = prev * (1.0 - min_margin/100.0)
    chosen=None; chosen_key=None; chosen_margin=None
    priority = ["price_optimal", "price_aggressive", "price_base"]
    for key in priority:
        p = candidate_prices.get(key)
        if p is None: continue
        margin_pct = (p - cost) / (p + 1e-9) * 100.0
        if margin_pct >= min_margin:
            chosen = p; chosen_key = key; chosen_margin = margin_pct; break
    if chosen is None:
        best_m=-math.inf; best_k=None; best_p=None
        for k,v in candidate_prices.items():
            if v is None: continue
            m = (v - cost) / (v + 1e-9) * 100.0
            if m > best_m:
                best_m=m; best_k=k; best_p=v
        if best_p is None:
            best_p = float(row.get("local_price", 100.0)); best_k="fallback"
            best_m = (best_p - cost) / (best_p + 1e-9) * 100.0
        chosen = best_p; chosen_key = best_k; chosen_margin = best_m
    prev_price = float(row.get("local_price", chosen))
    change_pct = abs(pct_change(prev_price, chosen))
    max_monthly = guardrails.get("max_monthly_change_pct", 15.0)
    if change_pct > max_monthly:
        allowed = prev_price * (1 + math.copysign(max_monthly/100.0, chosen - prev_price))
        chosen = round(allowed, 2)
        change_pct = abs(pct_change(prev_price, chosen))
        chosen_margin = (chosen - cost) / (chosen + 1e-9) * 100.0
    if change_pct <= guardrails.get("auto_approve_pct", 2.0):
        approval = "Auto-approve"
    elif change_pct <= guardrails.get("manager_review_pct", 5.0):
        approval = "Manager Review"
    else:
        approval = "Committee Review"
    return {"price_recommended": round(chosen,2), "source_key": chosen_key, "margin_pct": round(chosen_margin,2), "change_pct": round(change_pct,2), "approval": approval}

# ---------------- Printing helpers ----------------
def print_agent_traces(traces):
    console.print(Panel("[bold]Agent outputs (3 candidates each)[/bold]"))
    for aname, at in traces.items():
        cands = {k: format_inr(v) for k,v in at.candidates.items()}
        console.print(f"{aname.upper():<12} Candidates: {json.dumps(cands)}  Details: {json.dumps(at.details)}", markup=False)

def print_candidates_table(candidates_detail, candidate_explanations, selection):
    console.print(Panel("[bold]Candidates & Predictions[/bold]"))
    table = Table(show_lines=True)
    table.add_column("Candidate")
    table.add_column("Price", justify="right")
    table.add_column("Pred Units", justify="right")
    table.add_column("Elasticity", justify="right")
    table.add_column("LLM Explanation", overflow="fold")
    for key, info in candidates_detail.items():
        expl = candidate_explanations.get(key, "")
        table.add_row(key, format_inr(info["price"]), f"{info['pred_units']:.1f}", f"{info['elasticity']:.2f}", expl[:400])
    console.print(table)
    console.print(Panel("[bold]Selection[/bold]\n" + json.dumps(selection, indent=2)))

# ---------------- Light NL parser for What-If (IMPROVED) ----------------
def parse_nl_scenario(text: str) -> Dict[str, Any]:
    """
    Improved NL parser:
      - Detects local price, competitor price (including phrases like 'competitor drops to 290'),
      - Recognizes promo depth (-0.03 or -3% or 3%),
      - Inventory days,
      - Channel mentions and SKU mentions.
    Returns overrides dict (keys: local_price, competitor_price, promo_depth_pct, inventory_days, lead_time_days, purchase_cost, mrp, tax_pct, channel, sku_id).
    """
    text_raw = text
    text = text.lower()
    overrides = {}
    num_re = r"([0-9]+(?:\.[0-9]+)?)"

    # competitor patterns: 'competitor drops to 290', 'competitor to 290', 'competitor = 290', 'comp drops to 290'
    m = re.search(r"(?:competitor|comp|rival|peer)\s*(?:drops\s*to|down\s*to|to|=|is|becomes)\s*₹?\s*"+num_re, text)
    if m:
        overrides["competitor_price"] = float(m.group(1))

    # competitor patterns like 'drops by 10%' => relative
    m = re.search(r"(?:competitor|comp)\s*(?:drops|down)\s*(?:by)?\s*"+num_re+r"\s*%+", text)
    if m and "competitor_price" not in overrides:
        try:
            pct = float(m.group(1))
            # can't compute absolute price without baseline; leave a special flag to indicate relative competitor drop
            overrides["competitor_pct_change"] = -abs(pct)
        except Exception:
            pass

    # local price
    m = re.search(r"(?:local price|my price|our price|price)\s*(?:is|=|at|to|becomes)?\s*₹?\s*"+num_re, text)
    if m:
        overrides["local_price"] = float(m.group(1))

    # promo depth: 'promo -0.03', 'promo -3%', 'promo 3%'
    m = re.search(r"(?:promo|promotion|discount|promo depth)\s*(?:is|=|at)?\s*([+-]?[0-9]+(?:\.[0-9]+)?)\s*%?", text)
    if m:
        val = float(m.group(1))
        if abs(val) < 1 and "." in m.group(1):
            val = val * 100.0
        overrides["promo_depth_pct"] = float(val)

    # inventory days
    m = re.search(r"(?:inventory days|inventory)\s*(?:is|=|at|:)?\s*"+num_re, text)
    if m:
        overrides["inventory_days"] = float(m.group(1))

    m = re.search(r"(?:lead time|leadtime|supplier lead time)\s*(?:is|=|at|:)?\s*"+num_re, text)
    if m:
        overrides["lead_time_days"] = float(m.group(1))
    m = re.search(r"(?:purchase cost|cost price|procurement cost)\s*(?:is|=|at|:)?\s*â‚¹?\s*"+num_re, text)
    if m:
        overrides["purchase_cost"] = float(m.group(1))
    m = re.search(r"(?:mrp)\s*(?:is|=|at|:)?\s*â‚¹?\s*"+num_re, text)
    if m:
        overrides["mrp"] = float(m.group(1))
    m = re.search(r"(?:tax|gst|tax pct)\s*(?:is|=|at|:)?\s*"+num_re+r"\s*%?", text)
    if m:
        overrides["tax_pct"] = float(m.group(1))

    # channel
    if re.search(r"\becom\b|\be-commerce\b|\bonline\b", text):
        overrides["channel"] = "ECOM"
    elif re.search(r"\bgt\b|\bgeneral trade\b", text):
        overrides["channel"] = "GT"
    elif re.search(r"\bmt\b|\bmodern trade\b", text):
        overrides["channel"] = "MT"

    # SKU patterns (SKU_0001, GCPL_SKU_0001, sku 0001)
    m = re.search(r"(sku[_\-\s]?[0-9a-zA-Z]+|gcpl[_\-\s]?sku[_\-\s]?[0-9a-zA-Z]+)", text, flags=re.I)
    if m:
        overrides["sku_id"] = m.group(1).upper().replace(" ", "_")

    # fallback: detect plain numbers that appear near 'competitor' words with different phrasing
    if "competitor_price" not in overrides:
        m2 = re.search(r"competitor.*?([0-9]+(?:\.[0-9]+)?)", text)
        if m2:
            overrides["competitor_price"] = float(m2.group(1))

    # ensure percentages are within guardrails later; return overrides
    return overrides

def nl_interpret_with_bedrock(nl_text: str, context: Dict[str, Any]) -> Dict[str, Any]:
    """
    If Bedrock configured, ask the LLM to return a JSON with:
      {
        "overrides": { "local_price":..., "competitor_price":..., "promo_depth_pct":..., "inventory_days":..., "lead_time_days":..., "purchase_cost":..., "mrp":..., "tax_pct":..., "channel":... },
        "extraction_confidence": "low|medium|high",
        "notes": "short text"
      }
    If parsing fails or Bedrock not available, fallback to parse_nl_scenario.
    IMPORTANT: LLM is instructed to only output JSON exactly in the required format or a one-line error; we validate JSON strictly.
    """
    parsed = parse_nl_scenario(nl_text)
    if boto3:
        try:
            system = {
                "role": "system",
                "content": (
                    "You are a strict extractor. Given a short natural-language pricing scenario, "
                    "output *ONLY* a single JSON object (no surrounding text) with keys: overrides, extraction_confidence, notes. "
                    "The 'overrides' value must be a JSON object containing any of: local_price (number), competitor_price (number), "
                    "promo_depth_pct (number, percent), inventory_days (number), lead_time_days (number), purchase_cost (number), mrp (number), tax_pct (number), channel (ECOM|GT|MT), sku_id (string). "
                    "Do NOT propose business recommendations or new prices. If uncertain, set extraction_confidence to 'low' and include best-guess values in overrides."
                )
            }
            user = {
                "role": "user",
                "content": f"Scenario: {nl_text}\nContext (optional): {json.dumps(context)[:1200]}"
            }
            messages = [system, user]
            out = bedrock_invoke(messages, max_tokens=250, temperature=0.0)
            # try to find the first JSON object in the output
            jtxt = None
            try:
                # locate first '{' and last '}' that make a valid JSON
                first = out.find('{')
                last = out.rfind('}')
                if first != -1 and last != -1:
                    jtxt = out[first:last+1]
                    obj = json.loads(jtxt)
                    # validate shape
                    if "overrides" in obj and isinstance(obj["overrides"], dict):
                        # coerce numeric strings to numbers where possible
                        cleaned = {}
                        for k, v in obj["overrides"].items():
                            try:
                                if isinstance(v, (int, float)):
                                    cleaned[k] = float(v)
                                elif isinstance(v, str) and re.match(r"^-?\d+(\.\d+)?$", v.strip()):
                                    cleaned[k] = float(v.strip())
                                else:
                                    cleaned[k] = v
                            except Exception:
                                cleaned[k] = v
                        obj["overrides"] = cleaned
                        return obj
            except Exception:
                pass
            # if bedrock output invalid, fallback to parser
            return {"overrides": parsed, "extraction_confidence": "low", "notes": "Bedrock output invalid or could not parse; falling back to rule-based parser."}
        except Exception as e:
            console.print(f"Bedrock interpretation failed: {e}", markup=False)
            return {"overrides": parsed, "extraction_confidence": "low", "notes": f"Bedrock error: {e}. Used rule-based parser."}
    else:
        return {"overrides": parsed, "extraction_confidence": "medium", "notes": "Bedrock unavailable; used local parser."}

RUNTIME_CACHE: Dict[str, Dict[str, Any]] = {}

def get_dataset_catalog() -> Dict[str, str]:
    out = {}
    for k, v in DATASET_CATALOG.items():
        if os.path.exists(v):
            out[k] = v
    # Fallback to env-selected dataset if not in map.
    if not out and os.path.exists(DATA_CSV):
        out["default"] = DATA_CSV
    # Ensure at least one entry so init_runtime never gets empty catalog (e.g. on EC2 when CSVs missing).
    if not out:
        out["raju"] = str(DATA_CSV)
    return out

def init_runtime(force_reload: bool = False, dataset_key: str = "raju") -> Dict[str, Any]:
    dataset_key = str(dataset_key or "raju").strip().lower()
    catalog = get_dataset_catalog()
    if dataset_key not in catalog:
        dataset_key = next(iter(catalog.keys()))
    if (not force_reload) and dataset_key in RUNTIME_CACHE and RUNTIME_CACHE[dataset_key].get("df") is not None:
        return RUNTIME_CACHE[dataset_key]
    ensure_output_dir()
    sop = parse_sop_docx(SOP_DOCX)
    guardrails = sop.get("guardrails", DEFAULT_GUARDRAILS)
    df = load_and_enrich(catalog[dataset_key])
    nn_model, nn_scaler = None, None
    if ENABLE_NN:
        try:
            nn_model, nn_scaler = load_nn_model("nn_demand_model.pt", "nn_scaler.npz")
            console.print("NN demand model loaded.", markup=False)
        except Exception as e:
            console.print(f"NN model not loaded: {e}", markup=False)
    policy_model = None
    if ENABLE_RL:
        try:
            if os.path.exists("policy_rf.pkl"):
                policy_model = load_policy("policy_rf.pkl")
                console.print("Policy model loaded.", markup=False)
        except Exception as e:
            console.print(f"Policy load failed: {e}", markup=False)
    runtime = {
        "df": df,
        "sop": sop,
        "guardrails": guardrails,
        "nn_model": nn_model,
        "nn_scaler": nn_scaler,
        "policy_model": policy_model,
        "dataset_key": dataset_key,
        "dataset_path": catalog[dataset_key],
    }
    RUNTIME_CACHE[dataset_key] = runtime
    return runtime

def _build_synthetic_row(df: pd.DataFrame, sku: str = "", channel: str = "", region: str = "") -> pd.Series:
    row = pd.Series(dtype=object)
    row["sku_id"] = sku or "SYNTHETIC_SKU"
    row["channel"] = channel or df["channel"].mode().iloc[0]
    row["region"] = region or df["region"].mode().iloc[0]
    row["category"] = df["category"].mode().iloc[0]
    for col in [
        "local_price","competitor_price","promo_depth_pct","final_price","units_sold","inventory_level","revenue","margin",
        "festival_lift","seasonality_index","lead_time_days","purchase_cost","mrp","tax_pct",
        "recent_7d_avg_units","recent_28d_avg_units","demand_trend_7_28","is_weekend","is_festival_day","is_pre_festival_window","hour_start"
    ]:
        row[col] = float(df[col].median(skipna=True)) if col in df.columns else 0.0
    row["date"] = df["date"].max()
    return row

def pick_row(df: pd.DataFrame, sku: str, channel: str, region: str) -> pd.Series:
    sku = resolve_input_value(sku, sorted(df["sku_id"].astype(str).unique().tolist()), "SKU_ID")
    channel = resolve_input_value(channel, sorted(df["channel"].astype(str).unique().tolist()), "Channel")
    region = resolve_input_value(region, sorted(df["region"].astype(str).unique().tolist()), "Region")
    sort_cols = [c for c in ["date", "txn_timestamp"] if c in df.columns]
    mask = (
        (df["sku_id"].astype(str) == str(sku)) &
        (df["channel"].astype(str).str.lower() == str(channel).lower()) &
        (df["region"].astype(str).str.lower() == str(region).lower())
    )
    if mask.sum() > 0:
        scoped = df[mask].sort_values(sort_cols) if sort_cols else df[mask]
        return scoped.iloc[-1]
    sku_mask = (df["sku_id"].astype(str) == str(sku))
    if sku_mask.sum() > 0:
        scoped = df[sku_mask].sort_values(sort_cols) if sort_cols else df[sku_mask]
        return scoped.iloc[-1]
    return _build_synthetic_row(df, sku=sku, channel=channel, region=region)

def run_pricing_pipeline(row: pd.Series, runtime: Dict[str, Any]) -> Dict[str, Any]:
    df = runtime["df"]
    guardrails = runtime["guardrails"]
    nn_model = runtime["nn_model"]
    nn_scaler = runtime["nn_scaler"]
    b = BasePriceAgent().run(row, df)
    p = PromoAgent(guardrails).run(row, b.candidates)
    c = CompetitorAgent().run(row, p.candidates)
    i = InventoryAgent().run(row, c.candidates)
    pr = ProcurementAgent().run(row, i.candidates)
    bl = BillingAgent().run(row, pr.candidates)
    traces = {"base": b, "promo": p, "competitor": c, "inventory": i, "procurement": pr, "billing": bl}
    elasticity = estimate_elasticity(df, sku=row.get("sku_id"), category=row.get("category"))
    for t in traces.values():
        t.details["estimated_elasticity"] = elasticity
    mapped = CandidateAssembler(guardrails).assemble(traces)
    demand_prior = safe_float(row.get("recent_7d_avg_units"), None)
    candidates = {}
    for k, price in mapped.items():
        pred_units = predict_units_for_candidate(row, price, elasticity, nn_model, nn_scaler, demand_prior=demand_prior)
        candidates[k] = {"price": round(price, 2), "pred_units": float(max(0.0, min(pred_units, 1e7))), "elasticity": elasticity}
    selection = enrich_selection_payload(row, selector_and_sop(row, mapped, guardrails), candidates)
    return {"traces": traces, "candidates": candidates, "selection": selection}


def derive_reorder_point(row: pd.Series) -> float:
    explicit = safe_float(row.get("reorder_point"), float("nan"))
    if np.isfinite(explicit):
        return max(0.0, explicit)
    lead = max(1.0, safe_float(row.get("lead_time_days"), 3.0))
    recent_7 = safe_float(row.get("recent_7d_avg_units"), float("nan"))
    recent_28 = safe_float(row.get("recent_28d_avg_units"), float("nan"))
    if not np.isfinite(recent_7):
        recent_7 = safe_float(row.get("units_sold"), 0.0)
    if not np.isfinite(recent_28):
        recent_28 = recent_7
    demand_floor = max(recent_7 * lead * 1.15, recent_28 * 0.45)
    return float(max(0.0, demand_floor))


def enrich_selection_payload(row: pd.Series, selection: Dict[str, Any], candidates: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    enriched = dict(selection)
    current_price = safe_float(row.get("local_price"), safe_float(row.get("final_price"), enriched.get("price_recommended", 0.0)))
    market_price = safe_float(row.get("competitor_price"), current_price)
    stock_on_hand = safe_float(row.get("stock_on_hand"), safe_float(row.get("inventory_level"), 0.0))
    opening_stock = safe_float(row.get("opening_stock"), stock_on_hand)
    reorder_point = derive_reorder_point(row)
    recent_7 = safe_float(row.get("recent_7d_avg_units"), safe_float(row.get("units_sold"), 0.0))
    recent_28 = safe_float(row.get("recent_28d_avg_units"), recent_7)
    inventory_days = stock_on_hand / max(recent_7, 1e-9)
    purchase_cost = safe_float(row.get("purchase_cost"), safe_float(row.get("cost"), 0.0))
    tax_pct = safe_float(row.get("tax_pct"), 0.0)
    recommended_price = safe_float(enriched.get("price_recommended"), current_price)
    unit_profit = recommended_price - purchase_cost
    price_gap_pct = pct_change(market_price, recommended_price)
    enriched.update({
        "sku_id": str(row.get("sku_id", "")),
        "item_name": str(row.get("item_name", "")),
        "category": str(row.get("category", "")),
        "supplier_id": str(row.get("supplier_id", "")),
        "festival_name": str(row.get("festival_name", "")),
        "price_current": round(current_price, 2),
        "market_price": round(market_price, 2),
        "price_gap_pct": round(price_gap_pct, 2),
        "promo_depth_pct": round(safe_float(row.get("promo_depth_pct"), 0.0), 2),
        "opening_stock": round(opening_stock, 2),
        "stock_on_hand": round(stock_on_hand, 2),
        "reorder_point": round(reorder_point, 2),
        "inventory_days_cover": round(inventory_days, 2),
        "needs_reorder": int(stock_on_hand <= reorder_point),
        "lead_time_days": round(safe_float(row.get("lead_time_days"), 0.0), 1),
        "purchase_cost": round(purchase_cost, 2),
        "mrp": round(safe_float(row.get("mrp"), recommended_price), 2),
        "tax_pct": round(tax_pct, 2),
        "recent_7d_avg_units": round(recent_7, 2),
        "recent_28d_avg_units": round(recent_28, 2),
        "unit_profit_est": round(unit_profit, 2),
        "net_profit_est": round(unit_profit * max(recent_7, 0.0), 2),
        "is_festival_day": int(safe_float(row.get("is_festival_day"), 0.0)),
        "is_pre_festival_window": int(safe_float(row.get("is_pre_festival_window"), 0.0)),
    })
    if candidates:
        try:
            enriched["candidate_prices"] = {k: round(safe_float(v.get("price"), 0.0), 2) for k, v in candidates.items()}
        except Exception:
            pass
    return enriched


def deterministic_price_explanation(selection: Dict[str, Any], candidates: Dict[str, Any]) -> str:
    item = selection.get("item_name") or selection.get("sku_id") or "this SKU"
    rec = safe_float(selection.get("price_recommended"), 0.0)
    market = safe_float(selection.get("market_price"), rec)
    margin = safe_float(selection.get("margin_pct"), 0.0)
    stock = safe_float(selection.get("stock_on_hand"), 0.0)
    reorder = safe_float(selection.get("reorder_point"), 0.0)
    gap = safe_float(selection.get("price_gap_pct"), 0.0)
    approval = selection.get("approval", "Pending")
    if stock <= reorder:
        stock_line = "Stock is close to or below reorder level, so price should stay disciplined until replenishment is secure."
    else:
        stock_line = "Stock cover is comfortable, so the engine can optimize for margin without creating shelf risk."
    gap_line = "close to the market" if abs(gap) <= 3 else ("above the market" if gap > 0 else "below the market")
    return (
        f"For {item}, the engine recommends {format_inr(rec)}. "
        f"This keeps margin near {margin:.1f}% and places the shelf price {gap_line} benchmark pricing "
        f"around {format_inr(market)}. {stock_line} Approval path: {approval}."
    )


def deterministic_whatif_explanation(original: Dict[str, Any], updated: Dict[str, Any], overrides: Dict[str, Any], scenario: str = "") -> str:
    orig_sel = original.get("selection", {})
    upd_sel = updated.get("selection", {})
    old_price = safe_float(orig_sel.get("price_recommended"), 0.0)
    new_price = safe_float(upd_sel.get("price_recommended"), 0.0)
    old_margin = safe_float(orig_sel.get("margin_pct"), 0.0)
    new_margin = safe_float(upd_sel.get("margin_pct"), 0.0)
    item = upd_sel.get("item_name") or upd_sel.get("sku_id") or "this SKU"
    key_moves = []
    if "competitor_price" in overrides:
        key_moves.append(f"competitor moved to {format_inr(overrides['competitor_price'])}")
    if "inventory_days" in overrides:
        key_moves.append(f"stock cover changed to {safe_float(overrides['inventory_days'], 0.0):.1f} days")
    if "promo_depth_pct" in overrides:
        key_moves.append(f"promo depth changed to {safe_float(overrides['promo_depth_pct'], 0.0):.1f}%")
    scenario_line = "; ".join(key_moves) if key_moves else "the scenario assumptions"
    direction = "up" if new_price >= old_price else "down"
    return (
        f"Under this scenario for {item}, the recommendation moves {direction} from {format_inr(old_price)} to {format_inr(new_price)}. "
        f"The main driver is {scenario_line}. Expected margin shifts from {old_margin:.1f}% to {new_margin:.1f}%. "
        f"Use this output to explain why the store should hold, react, or defend share."
    )


def deterministic_forecast_explanation(first_row: Dict[str, Any]) -> str:
    sel = first_row.get("selection", {}) or {}
    dq = first_row.get("demand_quantiles", {}) or {}
    item = sel.get("item_name") or sel.get("sku_id") or "this SKU"
    p50 = safe_float(dq.get("p50"), 0.0)
    p90 = safe_float(dq.get("p90"), p50)
    stock = safe_float(sel.get("stock_on_hand"), 0.0)
    reorder = safe_float(sel.get("reorder_point"), 0.0)
    price = safe_float(sel.get("price_recommended"), 0.0)
    if stock <= reorder:
        risk_line = "Current stock is at or below reorder level, so replenishment should happen immediately."
    elif stock < p50:
        risk_line = "Current stock is below expected demand, so watch daily sell-through closely."
    else:
        risk_line = "Current stock appears sufficient for the near-term forecast window."
    return (
        f"For {item}, expected demand is about {p50:.1f} units on a normal outcome and may reach {p90:.1f} units on a high-demand day. "
        f"Recommended price for the first forecast day is {format_inr(price)}. {risk_line}"
    )


def deterministic_detail_explanation(mode: str, payload: Dict[str, Any]) -> str:
    if mode == "price":
        sel = payload.get("selection", {}) or {}
        item = sel.get("item_name") or sel.get("sku_id") or "the selected SKU"
        return (
            f"Executive Summary\n"
            f"{item} is recommended at {format_inr(sel.get('price_recommended', 0))}. "
            f"This keeps the business inside margin discipline while staying grounded in current market pricing.\n\n"
            f"Key Numbers\n"
            f"Current shelf price: {format_inr(sel.get('price_current', 0))}\n"
            f"Recommended price: {format_inr(sel.get('price_recommended', 0))}\n"
            f"Market reference: {format_inr(sel.get('market_price', 0))}\n"
            f"Expected margin: {safe_float(sel.get('margin_pct'), 0.0):.1f}%\n"
            f"Inventory cover: {safe_float(sel.get('inventory_days_cover'), 0.0):.1f} days\n\n"
            f"Business Rationale\n"
            f"The engine balances three things together: margin, competitive position, and stock safety. "
            f"This avoids a price move that looks attractive on paper but hurts sell-through or stock availability in store.\n\n"
            f"Decision Guidance\n"
            f"Approval path: {sel.get('approval', 'Pending')}. "
            f"If market prices stay stable, this recommendation can be used as the operating shelf price for the next review cycle."
        )
    if mode == "whatif":
        return (
            f"Executive Summary\n"
            f"The engine has re-evaluated the price based on the scenario entered by the user.\n\n"
            f"Scenario Input\n"
            f"{payload.get('scenario', 'No scenario text provided.')}\n\n"
            f"Parsed Business Changes\n"
            f"{json.dumps(payload.get('overrides', {}), indent=2)}\n\n"
            f"Business Rationale\n"
            f"The recommendation changes only when the scenario materially affects competitiveness, stock cover, or promotion pressure. "
            f"This helps the owner see whether the situation needs a defensive reaction, a hold decision, or a margin-first response.\n\n"
            f"Decision Guidance\n"
            f"Use the before-versus-after block to explain the commercial trade-off, and use the candidate deltas to show how strongly each pricing path reacted."
        )
    first = payload.get("forecast_first_row", {}) or {}
    sel = first.get("selection", {}) or {}
    dq = first.get("demand_quantiles", {}) or {}
    return (
        f"Executive Summary\n"
        f"The forecast gives a planning view for the next selling window and combines demand expectation with inventory risk.\n\n"
        f"Key Numbers\n"
        f"Forecast start: {first.get('date', '')}\n"
        f"Item: {sel.get('item_name') or sel.get('sku_id') or 'Selected SKU'}\n"
        f"Base demand case (P50): {safe_float(dq.get('p50'), 0.0):.1f}\n"
        f"High demand case (P90): {safe_float(dq.get('p90'), 0.0):.1f}\n"
        f"Recommended price: {format_inr(sel.get('price_recommended', 0))}\n"
        f"Stock on hand: {safe_float(sel.get('stock_on_hand'), 0.0):.1f}\n"
        f"Reorder point: {safe_float(sel.get('reorder_point'), 0.0):.1f}\n\n"
        f"Business Rationale\n"
        f"P50 should be treated as the working plan and P90 as the stress case. "
        f"If stock is near or below reorder point, replenishment must move ahead of promotional expansion.\n\n"
        f"Decision Guidance\n"
        f"This view is designed to support one combined decision: how much to stock, what price to hold, and whether the store is exposed to a demand spike."
    )


def sanitize_llm_explanation(text: Any) -> str:
    raw = str(text or "").strip()
    if not raw:
        return ""
    cleaned = re.sub(r"<reasoning>.*?</reasoning>", "", raw, flags=re.IGNORECASE | re.DOTALL)
    cleaned = re.sub(r"</?(thinking|analysis|scratchpad|reflection)>.*?</\1>", "", cleaned, flags=re.IGNORECASE | re.DOTALL)
    cleaned = cleaned.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
    cleaned = re.sub(r"<[^>]+>", "", cleaned)
    cleaned = cleaned.replace("**", "").replace("__", "")
    cleaned = re.sub(r"^#{1,6}\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^-{3,}$", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"-{4,}", " - ", cleaned)
    cleaned = re.sub(r"^\|.*\|$", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^[\s|\-:]{3,}$", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^USER:\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^ASSISTANT:\s*", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"^Section\s*\|\s*Key Points\s*$", "", cleaned, flags=re.MULTILINE | re.IGNORECASE)
    cleaned = re.sub(r"\bwe need to produce\b.*?(?=\n|$)", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\blet'?s craft\b.*?(?=\n|$)", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\bprovide sections\b.*?(?=\n|$)", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\bkeep it readable\b.*?(?=\n|$)", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\buse simple business language\b.*?(?=\n|$)", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"[ \t]*\|[ \t]*", " - ", cleaned)
    cleaned = re.sub(r"^[ \t-]+$", "", cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def summarize_explanation_text(text: str, max_chars: int = 320) -> str:
    cleaned = sanitize_llm_explanation(text)
    if not cleaned:
        return ""
    first_block = next((part.strip() for part in re.split(r"\n{2,}", cleaned) if part.strip()), cleaned)
    if len(first_block) <= max_chars:
        return first_block
    cut = first_block[:max_chars].rsplit(" ", 1)[0].strip()
    return f"{cut}..."


def ensure_full_explanation_text(detail_text: str, fallback_text: str) -> str:
    cleaned = sanitize_llm_explanation(detail_text)
    if cleaned:
        return cleaned
    return sanitize_llm_explanation(fallback_text)


def resolve_explanation_bundle(
    explanation: Dict[str, str],
    explanation_detail: Dict[str, str],
    fallback_detail: str,
) -> Dict[str, str]:
    short_text = sanitize_llm_explanation(explanation.get("text", ""))
    detail_text = sanitize_llm_explanation(explanation_detail.get("text", ""))
    source = "bedrock" if explanation.get("source") == "bedrock" or explanation_detail.get("source") == "bedrock" else explanation.get("source", "deterministic")
    error = explanation.get("error") or explanation_detail.get("error") or ""
    if REQUIRE_BEDROCK:
        if not detail_text or source != "bedrock":
            raise RuntimeError(error or "Bedrock explanation is required but was not returned by the backend.")
    else:
        detail_text = detail_text or sanitize_llm_explanation(fallback_detail)
    short_text = short_text or summarize_explanation_text(detail_text)
    return {"short_text": short_text, "detail_text": detail_text, "source": source, "error": error}


def build_business_prompt(task: str, payload: Dict[str, Any], detail: bool = False) -> str:
    sections = (
        "1. Executive Summary\n"
        "2. Key Numbers\n"
        "3. Business Rationale\n"
        "4. Risks To Watch\n"
        "5. Recommended Action"
        if detail else
        "1. Decision\n2. Why It Makes Sense\n3. What To Do Next"
    )
    return (
        "You are AI Sahayak, writing for a small retail owner and a hackathon judging panel.\n"
        "Return plain text only.\n"
        "Do not include XML tags, HTML tags, markdown tables, code fences, bold markers, or hidden reasoning.\n"
        "Do not reveal chain-of-thought.\n"
        "Never write meta lines such as 'we need to produce', 'let's craft', 'parsed input table', or anything about how you are writing the answer.\n"
        "Use confident, business-friendly language.\n"
        "Keep every number aligned with the supplied engine output.\n"
        "Use short section headings and compact bullet-style sentences where useful.\n"
        "Write as if presenting to an owner who wants a commercial decision note, not a model transcript.\n"
        f"Task: {task}\n"
        f"Required sections:\n{sections}\n\n"
        f"Engine payload:\n{json.dumps(payload, indent=2, ensure_ascii=True)}"
    )


def _looks_like_meta_explanation(text: str) -> bool:
    sample = (text or "").lower()
    bad_markers = [
        "<reasoning>",
        "we need to produce",
        "let's craft",
        "provide sections",
        "parsed input",
        "| section |",
        "| what it means |",
        "<br>",
    ]
    return any(marker in sample for marker in bad_markers)


def _looks_incomplete_explanation(text: str) -> bool:
    sample = (text or "").strip()
    if not sample:
        return True
    if len(sample) < 120:
        return False
    lowered = sample.lower()
    if lowered.endswith(("because", "therefore", "however", "while", "and", "or", "with", "but")):
        return True
    if sample.endswith((",", ":", ";", "-", "(", "/")):
        return True
    if not re.search(r"[.!?]\"?$", sample):
        return True
    return False


def _rewrite_as_business_note(text: str, max_tokens: int) -> str:
    rewritten = bedrock_invoke(
        [
            {
                "role": "system",
                "content": (
                    "Rewrite the input into a polished retail business note. "
                    "Return plain text only. No XML, no markdown table, no HTML, no chain-of-thought, no meta commentary."
                ),
            },
            {
                "role": "user",
                "content": (
                    "Rewrite this into clean business English with these sections: "
                    "Executive Summary, Key Numbers, Business Rationale, Risks To Watch, Recommended Action.\n\n"
                    f"Input:\n{text}"
                ),
            },
        ],
        max_tokens=max_tokens,
        temperature=0.1,
    )
    return sanitize_llm_explanation(rewritten)


def maybe_bedrock_explanation(prompt: str, fallback_text: str, max_tokens: int = 180) -> Dict[str, str]:
    try:
        if boto3 is None:
            if REQUIRE_BEDROCK:
                return {"text": "", "source": "bedrock_error", "error": "boto3 is not installed on the backend."}
            return {"text": sanitize_llm_explanation(fallback_text), "source": "deterministic"}
        bedrock_status = get_bedrock_status()
        if REQUIRE_BEDROCK and not bedrock_status.get("ok"):
            return {"text": "", "source": "bedrock_error", "error": bedrock_status.get("error", "Bedrock is not ready.")}
        text = bedrock_invoke(
            [
                {"role": "system", "content": "You are AI Sahayak. Produce polished retail-business explanations in plain text only. Never output XML tags, HTML, markdown tables, code fences, or chain-of-thought."},
                {"role": "user", "content": prompt},
            ],
            max_tokens=max_tokens,
            temperature=0.2,
        )
        cleaned = sanitize_llm_explanation(text)
        if cleaned and _looks_like_meta_explanation(cleaned):
            cleaned = _rewrite_as_business_note(cleaned, max_tokens=max(max_tokens, 700))
        if cleaned and _looks_incomplete_explanation(cleaned):
            cleaned = _rewrite_as_business_note(cleaned, max_tokens=max(max_tokens, 700))
        return {"text": cleaned or sanitize_llm_explanation(fallback_text), "source": "bedrock" if cleaned else "deterministic"}
    except Exception as exc:
        if REQUIRE_BEDROCK:
            return {"text": "", "source": "bedrock_error", "error": str(exc)}
        return {"text": sanitize_llm_explanation(fallback_text), "source": "deterministic"}

def compute_kpis(df: pd.DataFrame, sku_id: Optional[str] = None) -> Dict[str, Any]:
    d = df.copy()
    if sku_id:
        m = d["sku_id"].astype(str) == str(sku_id)
        if m.sum() > 0:
            d = d[m]
    if d.empty:
        return {"kpis": {}, "series": {}}
    d = d.sort_values("date")
    daily = d.groupby("date", as_index=False).agg(
        revenue=("revenue", "sum"),
        profit=("profit_amt", "sum"),
        net_profit=("net_profit", "sum"),
        units=("units_sold", "sum"),
        stockout=("is_stockout", "mean"),
        promo=("promo_depth_pct", "mean"),
        margin=("margin", "mean"),
        price=("local_price", "mean"),
        market_price=("competitor_price", "mean"),
        tax_pct=("tax_pct", "mean"),
        reorder=("inventory_level", "mean"),
        festival=("is_festival_day", "max"),
    )
    daily["price_gap_pct"] = ((daily["price"] - daily["market_price"]) / (daily["market_price"].replace(0, np.nan) + 1e-9) * 100.0).replace([np.inf, -np.inf], np.nan).fillna(0.0)
    last7 = daily.tail(7)
    last30 = daily.tail(30)
    prev30 = daily.iloc[max(0, len(daily) - 60):max(0, len(daily) - 30)]

    def pct_delta(a, b):
        return float((a - b) / (abs(b) + 1e-9) * 100.0)

    rev30 = float(last30["revenue"].sum())
    prof30 = float(last30["profit"].sum())
    units30 = float(last30["units"].sum())
    avg_price30 = float(last30["price"].mean())
    avg_margin30 = float(last30["margin"].mean())
    avg_stockout30 = float(last30["stockout"].mean() * 100.0)
    promo30 = float(last30["promo"].mean())
    rev_per_unit = float(rev30 / (units30 + 1e-9))
    inv_days = float((last30["reorder"].mean() / (last30["units"].mean() / 30.0 + 1e-9)))
    fest_days_30 = int(last30["festival"].sum())
    avg_gap30 = float(last30["price_gap_pct"].mean())
    net_profit30 = float(last30["net_profit"].sum())
    avg_tax30 = float(last30["tax_pct"].mean())
    sell_through_30 = float((last30["units"].sum() / (last30["reorder"].sum() + last30["units"].sum() + 1e-9)) * 100.0)

    rev_prev = float(prev30["revenue"].sum()) if not prev30.empty else rev30
    prof_prev = float(prev30["profit"].sum()) if not prev30.empty else prof30
    units_prev = float(prev30["units"].sum()) if not prev30.empty else units30

    latest_rows = d.sort_values("date").groupby("sku_id", as_index=False).tail(1).copy()
    latest_rows["reorder_point_est"] = latest_rows.apply(derive_reorder_point, axis=1)
    latest_rows["inventory_days_cover"] = latest_rows["inventory_level"] / (latest_rows["recent_7d_avg_units"].replace(0, np.nan) + 1e-9)
    latest_rows["reorder_risk"] = (latest_rows["inventory_level"] <= latest_rows["reorder_point_est"]).astype(int)
    reorder_risk_skus = int(latest_rows["reorder_risk"].sum())
    low_cover_skus = int((latest_rows["inventory_days_cover"].fillna(0) < latest_rows["lead_time_days"].fillna(3)).sum())

    top_skus = (
        d.groupby(["sku_id", "item_name"], as_index=False)
        .agg(revenue=("revenue", "sum"), profit=("profit_amt", "sum"), units=("units_sold", "sum"), stock=("inventory_level", "last"))
        .sort_values("revenue", ascending=False)
        .head(6)
    )
    category_mix = (
        d.groupby("category", as_index=False)
        .agg(revenue=("revenue", "sum"), units=("units_sold", "sum"), profit=("profit_amt", "sum"))
        .sort_values("revenue", ascending=False)
        .head(6)
    )
    payment_mix = []
    if "payment_mode" in d.columns:
        payment_mix = (
            d.groupby("payment_mode", as_index=False)
            .agg(revenue=("revenue", "sum"), txns=("sku_id", "count"))
            .sort_values("revenue", ascending=False)
            .to_dict(orient="records")
        )

    alerts = []
    if reorder_risk_skus > 0:
        alerts.append(f"{reorder_risk_skus} SKUs are below reorder point.")
    if avg_gap30 > 5:
        alerts.append("Your average price is above market; watch conversion risk.")
    elif avg_gap30 < -5:
        alerts.append("You are priced below market on average; margin opportunity exists.")
    if promo30 > 10:
        alerts.append("Promotions are running deep; review margin leakage.")
    if fest_days_30 > 0:
        alerts.append(f"{fest_days_30} festival-linked day(s) influenced the recent demand window.")

    kpis = {
        "revenue_30d": rev30,
        "profit_30d": prof30,
        "net_profit_30d": net_profit30,
        "units_30d": units30,
        "avg_price": avg_price30,
        "avg_margin_pct": avg_margin30,
        "stockout_pct": avg_stockout30,
        "avg_promo_depth_pct": promo30,
        "revenue_per_unit": rev_per_unit,
        "inventory_days_est": inv_days,
        "festival_days_last30": fest_days_30,
        "avg_price_gap_pct": avg_gap30,
        "avg_tax_pct": avg_tax30,
        "sell_through_pct": sell_through_30,
        "reorder_risk_skus": reorder_risk_skus,
        "low_cover_skus": low_cover_skus,
        "revenue_growth_pct": pct_delta(rev30, rev_prev),
        "profit_growth_pct": pct_delta(prof30, prof_prev),
        "units_growth_pct": pct_delta(units30, units_prev),
    }
    series = {
        "dates": [str(x.date()) if hasattr(x, "date") else str(x) for x in last30["date"].tolist()],
        "revenue": [float(x) for x in last30["revenue"].tolist()],
        "profit": [float(x) for x in last30["profit"].tolist()],
        "net_profit": [float(x) for x in last30["net_profit"].tolist()],
        "units": [float(x) for x in last30["units"].tolist()],
        "price": [float(x) for x in last30["price"].tolist()],
        "stockout_pct": [float(x * 100.0) for x in last30["stockout"].tolist()],
        "promo_depth_pct": [float(x) for x in last30["promo"].tolist()],
        "price_gap_pct": [float(x) for x in last30["price_gap_pct"].tolist()],
    }
    return {
        "kpis": kpis,
        "series": series,
        "top_skus": top_skus.to_dict(orient="records"),
        "category_mix": category_mix.to_dict(orient="records"),
        "payment_mix": payment_mix,
        "alerts": alerts,
    }

def create_api_app() -> Any:
    if Flask is None:
        raise RuntimeError("Flask is not installed. Install with: pip install flask")
    app = Flask(__name__)

    @app.after_request
    def _cors(resp):
        resp.headers["Access-Control-Allow-Origin"] = "*"
        resp.headers["Access-Control-Allow-Headers"] = "Content-Type, Authorization"
        resp.headers["Access-Control-Allow-Methods"] = "GET, POST, OPTIONS"
        return resp

    # Chat drive-ui: register first so they are always available (WP chat drives dashboard)
    @app.route("/api/drive-ui", methods=["POST", "OPTIONS"], strict_slashes=False)
    def drive_ui():
        """Called by agents backend when user asks for review/price in WP chat."""
        if request.method == "OPTIONS":
            return "", 204
        global _CHAT_DRIVE_STATE
        data = request.get_json(silent=True) or {}
        action = (data.get("action") or "").strip().lower() or "review"
        payload = data.get("payload")
        _CHAT_DRIVE_STATE = {
            "action": action if action in ("review", "price", "insights", "overview") else "review",
            "payload": payload,
            "ts": time.time(),
        }
        return jsonify({"ok": True, "action": _CHAT_DRIVE_STATE["action"]})

    @app.route("/api/chat-action", methods=["GET"], strict_slashes=False)
    def chat_action():
        """Polled by dashboard frontend. Returns last chat-driven action and payload; clears after read."""
        global _CHAT_DRIVE_STATE
        now = time.time()
        prev = _CHAT_DRIVE_STATE
        if prev and (now - prev.get("ts", 0)) < _CHAT_DRIVE_TTL_SEC:
            _CHAT_DRIVE_STATE = {}
            return jsonify({"action": prev.get("action", "review"), "payload": prev.get("payload")})
        return jsonify({"action": None, "payload": None})

    @app.route("/api/health", methods=["GET"])
    def health():
        return jsonify({"ok": True, "service": "ai-sahayak-api"})

    @app.route("/api/meta", methods=["GET"])
    def meta():
        try:
            dataset_key = request.args.get("dataset_key", "raju")
            rt = init_runtime(dataset_key=dataset_key)
            df = rt["df"]
            skus = (
                df.groupby("sku_id", as_index=False)
                .agg(item_name=("item_name", "first"), category=("category", "first"))
                .sort_values("sku_id")
            )
            return jsonify({
                "active_dataset": rt.get("dataset_key", dataset_key),
                "datasets": [{"key": k, "path": v} for k, v in get_dataset_catalog().items()],
                "sku_count": int(df["sku_id"].nunique()),
                "rows": int(len(df)),
                "date_min": str(df["date"].min().date()),
                "date_max": str(df["date"].max().date()),
                "skus": skus.to_dict(orient="records"),
                "channels": sorted(df["channel"].astype(str).unique().tolist()),
                "regions": sorted(df["region"].astype(str).unique().tolist()),
            })
        except Exception as e:
            return jsonify({
                "active_dataset": "raju",
                "datasets": [{"key": "raju", "path": str(DATA_CSV)}],
                "sku_count": 0,
                "rows": 0,
                "date_min": "",
                "date_max": "",
                "skus": [],
                "channels": ["GT"],
                "regions": ["West"],
            }), 200

    @app.route("/api/model-status", methods=["GET"])
    def model_status():
        try:
            dataset_key = request.args.get("dataset_key", "raju")
            rt = init_runtime(dataset_key=dataset_key)
            bedrock_status = get_bedrock_status(force=True)
            deepar_endpoint = get_deepar_endpoint(dataset_key)
            status = {
                "dataset_key": rt.get("dataset_key"),
                "dnn_loaded": rt.get("nn_model") is not None and rt.get("nn_scaler") is not None,
                "policy_loaded": rt.get("policy_model") is not None,
                "bedrock_ready": bool(bedrock_status.get("ok")),
                "bedrock_required": REQUIRE_BEDROCK,
                "bedrock_error": bedrock_status.get("error", ""),
                "bedrock_model_primary": BEDROCK_MODEL_ID,
                "bedrock_model_fallbacks": BEDROCK_FALLBACK_MODELS,
                "deepar_endpoint_configured": bool(deepar_endpoint),
                "forecast_primary": "DeepAR (SageMaker)" if deepar_endpoint else "DeepAR Proxy (local)",
                "aws_region": AWS_REGION,
                "deepar_endpoint": deepar_endpoint if deepar_endpoint else "",
            }
            return jsonify(status)
        except Exception as e:
            return jsonify({
                "dataset_key": "raju",
                "dnn_loaded": False,
                "policy_loaded": False,
                "bedrock_ready": False,
                "bedrock_required": REQUIRE_BEDROCK,
                "bedrock_error": str(e),
                "bedrock_model_primary": BEDROCK_MODEL_ID,
                "bedrock_model_fallbacks": BEDROCK_FALLBACK_MODELS,
                "deepar_endpoint_configured": False,
                "forecast_primary": "DeepAR Proxy (local)",
                "aws_region": AWS_REGION,
                "deepar_endpoint": "",
            }), 200

    @app.route("/api/kpis", methods=["GET"])
    def api_kpis():
        dataset_key = request.args.get("dataset_key", "raju")
        sku_id = request.args.get("sku_id", "")
        rt = init_runtime(dataset_key=dataset_key)
        out = compute_kpis(rt["df"], sku_id=sku_id if sku_id else None)
        return jsonify({"dataset_key": rt.get("dataset_key"), **out})

    @app.route("/api/price", methods=["POST"])
    def api_price():
        data = request.get_json(silent=True) or {}
        dataset_key = str(data.get("dataset_key", "raju"))
        rt = init_runtime(dataset_key=dataset_key)
        sku = str(data.get("sku_id", "")).strip()
        channel = str(data.get("channel", "GT")).strip()
        region = str(data.get("region", "West")).strip()
        row = pick_row(rt["df"], sku, channel, region).copy()
        for k in ("local_price", "competitor_price", "promo_depth_pct", "inventory_days", "lead_time_days", "purchase_cost", "mrp", "tax_pct"):
            if k in data and data[k] is not None and str(data[k]).strip() != "":
                row[k] = safe_float(data[k], row.get(k))
        res = run_pricing_pipeline(row, rt)
        explanation = maybe_bedrock_explanation(
            prompt=build_business_prompt("Explain the price recommendation for a business user.", {
                "selection": res.get("selection", {}),
                "candidates": res.get("candidates", {}),
            }, detail=False),
            fallback_text=deterministic_price_explanation(res.get("selection", {}), res.get("candidates", {})),
        )
        explanation_detail = maybe_bedrock_explanation(
            prompt=build_business_prompt("Write a decision memo for the pricing recommendation.", {
                "selection": res.get("selection", {}),
                "candidates": res.get("candidates", {}),
                "traces": to_serializable_traces(res.get("traces", {})),
            }, detail=True),
            fallback_text=deterministic_detail_explanation("price", {"selection": res.get("selection", {})}),
            max_tokens=1000,
        )
        fallback_detail_text = deterministic_detail_explanation("price", {"selection": res.get("selection", {})})
        try:
            resolved = resolve_explanation_bundle(explanation, explanation_detail, fallback_detail_text)
        except RuntimeError as exc:
            return jsonify({
                "ok": False,
                "error": str(exc),
                "assistant_source": "bedrock_error",
                "bedrock_required": REQUIRE_BEDROCK,
                "bedrock_status": get_bedrock_status(),
            }), 503
        out = serialize_pipeline_result({
            "sku_id": row.get("sku_id"),
            "item_name": row.get("item_name", ""),
            "category": row.get("category", ""),
            "dataset_key": rt.get("dataset_key"),
            "inputs": data,
            "assistant_message": resolved["short_text"],
            "assistant_source": resolved["source"],
            "assistant_detail": resolved["detail_text"],
            **res,
        })
        return jsonify(out)

    @app.route("/api/whatif", methods=["POST"])
    def api_whatif():
        data = request.get_json(silent=True) or {}
        dataset_key = str(data.get("dataset_key", "raju"))
        rt = init_runtime(dataset_key=dataset_key)
        sku = str(data.get("sku_id", "")).strip()
        channel = str(data.get("channel", "GT")).strip()
        region = str(data.get("region", "West")).strip()
        base_row = pick_row(rt["df"], sku, channel, region).copy()
        original = run_pricing_pipeline(base_row, rt)
        scenario = data.get("scenario")
        overrides = data.get("overrides", {})
        if scenario:
            parsed = nl_interpret_with_bedrock(str(scenario), {"sku": sku, "channel": channel, "region": region})
            if isinstance(parsed, dict):
                overrides = {**(parsed.get("overrides", {}) if isinstance(parsed.get("overrides", {}), dict) else {}), **overrides}
        changed = base_row.copy()
        for k, v in overrides.items():
            if k in ("local_price","competitor_price","promo_depth_pct","purchase_cost","mrp","tax_pct","lead_time_days"):
                changed[k] = safe_float(v, changed.get(k))
            elif k == "inventory_days":
                changed["inventory_days"] = safe_float(v, changed.get("inventory_days"))
            elif k in ("channel","region","sku_id","category"):
                changed[k] = str(v)
        updated = run_pricing_pipeline(changed, rt)
        delta = {}
        for k in original["candidates"]:
            try:
                delta[k] = round(updated["candidates"][k]["price"] - original["candidates"][k]["price"], 2)
            except Exception:
                delta[k] = None
        explanation = maybe_bedrock_explanation(
            prompt=build_business_prompt("Explain the what-if scenario result in plain business language.", {
                "scenario": scenario,
                "overrides": overrides,
                "original_selection": original.get("selection", {}),
                "updated_selection": updated.get("selection", {}),
                "delta": delta,
            }, detail=False),
            fallback_text=deterministic_whatif_explanation(original, updated, overrides, scenario=str(scenario or "")),
        )
        explanation_detail = maybe_bedrock_explanation(
            prompt=build_business_prompt("Write a scenario analysis memo for the store owner.", {
                "scenario": scenario,
                "overrides": overrides,
                "original_selection": original.get("selection", {}),
                "updated_selection": updated.get("selection", {}),
                "delta": delta,
            }, detail=True),
            fallback_text=deterministic_detail_explanation("whatif", {"scenario": scenario, "overrides": overrides}),
            max_tokens=1000,
        )
        fallback_detail_text = deterministic_detail_explanation("whatif", {"scenario": scenario, "overrides": overrides})
        try:
            resolved = resolve_explanation_bundle(explanation, explanation_detail, fallback_detail_text)
        except RuntimeError as exc:
            return jsonify({
                "ok": False,
                "error": str(exc),
                "assistant_source": "bedrock_error",
                "bedrock_required": REQUIRE_BEDROCK,
                "bedrock_status": get_bedrock_status(),
            }), 503
        return jsonify(serialize_pipeline_result({
            "sku_id": base_row.get("sku_id"),
            "dataset_key": rt.get("dataset_key"),
            "overrides": overrides,
            "original": original,
            "updated": updated,
            "delta": delta,
            "assistant_message": resolved["short_text"],
            "assistant_source": resolved["source"],
            "assistant_detail": resolved["detail_text"],
        }))

    @app.route("/api/forecast", methods=["POST"])
    def api_forecast():
        data = request.get_json(silent=True) or {}
        dataset_key = str(data.get("dataset_key", "raju"))
        rt = init_runtime(dataset_key=dataset_key)
        sku = str(data.get("sku_id", "")).strip()
        channel = str(data.get("channel", "GT")).strip()
        region = str(data.get("region", "West")).strip()
        start_date = str(data.get("start_date", (pd.Timestamp.now() + pd.Timedelta(days=1)).strftime("%Y-%m-%d")))
        days = int(data.get("days", 14))
        days = max(1, min(days, 90))
        out = forecast_and_price(
            rt["df"], rt["sop"], sku, channel, region, start_date, days, rt["nn_model"], rt["nn_scaler"], rt["policy_model"],
            festival_context=data.get("festival_context"),
            dataset_key=dataset_key,
        )
        explanation = maybe_bedrock_explanation(
            prompt=build_business_prompt("Explain the near-term demand forecast for a business user.", {
                "forecast_first_row": out[0] if out else {},
            }, detail=False),
            fallback_text=deterministic_forecast_explanation(out[0] if out else {}),
        ) if out else {"text": "No forecast rows were generated.", "source": "deterministic"}
        explanation_detail = maybe_bedrock_explanation(
            prompt=build_business_prompt("Write a forecast decision note covering demand, price, inventory risk, and action.", {
                "forecast_first_row": out[0] if out else {},
            }, detail=True),
            fallback_text=deterministic_detail_explanation("forecast", {"forecast_first_row": out[0] if out else {}}),
            max_tokens=1000,
        ) if out else {"text": "No forecast rows were generated.", "source": "deterministic"}
        fallback_detail_text = deterministic_detail_explanation("forecast", {"forecast_first_row": out[0] if out else {}})
        try:
            resolved = resolve_explanation_bundle(explanation, explanation_detail, fallback_detail_text)
        except RuntimeError as exc:
            return jsonify({
                "ok": False,
                "error": str(exc),
                "assistant_source": "bedrock_error",
                "bedrock_required": REQUIRE_BEDROCK,
                "bedrock_status": get_bedrock_status(),
            }), 503
        return jsonify(serialize_pipeline_result({
            "dataset_key": rt.get("dataset_key"),
            "forecast": out,
            "assistant_message": resolved["short_text"],
            "assistant_source": resolved["source"],
            "assistant_detail": resolved["detail_text"],
        }))

    @app.route("/api/bedrock/test", methods=["POST"])
    def api_bedrock_test():
        data = request.get_json(silent=True) or {}
        prompt = str(data.get("prompt", "Reply with: Bedrock connected"))
        try:
            text = bedrock_invoke([{"role": "user", "content": prompt}], max_tokens=80, temperature=0.0)
            return jsonify({"ok": True, "response": text[:500]})
        except Exception as e:
            return jsonify({"ok": False, "error": str(e)}), 500

    return app

# ---------------- Interactive app ----------------
def run_interactive():
    ensure_output_dir()
    sop = parse_sop_docx(SOP_DOCX)
    guardrails = sop.get("guardrails", DEFAULT_GUARDRAILS)
    df = load_and_enrich(DATA_CSV)

    nn_model, nn_scaler = None, None
    if ENABLE_NN:
        try:
            nn_model, nn_scaler = load_nn_model("nn_demand_model.pt", "nn_scaler.npz")
            console.print("NN demand model loaded.", markup=False)
        except Exception as e:
            console.print(f"NN model not loaded: {e}", markup=False)
            nn_model, nn_scaler = None, None

    policy_model = None
    if ENABLE_RL:
        try:
            if os.path.exists("policy_rf.pkl"):
                policy_model = load_policy("policy_rf.pkl")
                console.print("Policy model loaded.", markup=False)
            else:
                console.print("Policy model not found; using deterministic assembler.", markup=False)
        except Exception as e:
            console.print(f"Policy load failed: {e}", markup=False)
            policy_model = None

    console.print("GCPL Pricing Hybrid - Interactive (full features enabled)", markup=False)

    while True:
        console.print("\nOptions: 1=Single SKU pricing, 2=What-If, 3=Batch run, 4=Train NN, 5=Train Policy, 6=Forecast future + price, 7=Exit")
        ch = input("Choose: ").strip()
        if ch == "7":
            console.print("Exiting."); break

        if ch == "1":
            sku = input("SKU_ID: ").strip()
            channel = input("Channel: ").strip()
            region = input("Region: ").strip()
            sku = resolve_input_value(sku, sorted(df["sku_id"].astype(str).unique().tolist()), "SKU_ID")
            channel = resolve_input_value(channel, sorted(df["channel"].astype(str).unique().tolist()), "Channel")
            region = resolve_input_value(region, sorted(df["region"].astype(str).unique().tolist()), "Region")
            mask_exact = (df["sku_id"].astype(str) == sku) & (df["channel"].astype(str).str.lower() == channel.lower()) & (df["region"].astype(str).str.lower() == region.lower())
            if mask_exact.sum() > 0:
                row = df[mask_exact].iloc[0]
            else:
                if sku:
                    alt = df[df["sku_id"].astype(str).str.contains(sku)]
                    if len(alt) > 0:
                        row = alt.iloc[0]
                    else:
                        console.print("No exact match. Using aggregated medians from dataset for synthetic row (not first row).", markup=False)
                        row = pd.Series()
                        row["sku_id"] = sku or "SYNTHETIC_SKU"
                        row["channel"] = channel or df["channel"].mode().iloc[0]
                        row["region"] = region or df["region"].mode().iloc[0]
                        row["category"] = df["category"].mode().iloc[0]
                        for col in ["local_price","competitor_price","promo_depth_pct","final_price","units_sold","inventory_level","revenue","margin","festival_lift","seasonality_index","lead_time_days","purchase_cost","mrp","tax_pct"]:
                            row[col] = float(df[col].median(skipna=True))
                        row["date"] = df["date"].max()
            base_tr = BasePriceAgent().run(row, df)
            promo_tr = PromoAgent(guardrails).run(row, base_tr.candidates)
            comp_tr = CompetitorAgent().run(row, promo_tr.candidates)
            inv_tr = InventoryAgent().run(row, comp_tr.candidates)
            proc_tr = ProcurementAgent().run(row, inv_tr.candidates)
            bill_tr = BillingAgent().run(row, proc_tr.candidates)
            traces = {"base": base_tr, "promo": promo_tr, "competitor": comp_tr, "inventory": inv_tr, "procurement": proc_tr, "billing": bill_tr}
            mapped = CandidateAssembler(guardrails).assemble(traces)
            elasticity = estimate_elasticity(df, sku=row.get("sku_id"), category=row.get("category"))
            for t in traces.values():
                t.details["estimated_elasticity"] = elasticity
            candidates_detail = {}
            for k, price in mapped.items():
                pred_units = predict_units_for_candidate(row, price, elasticity, nn_model, nn_scaler)
                candidates_detail[k] = {"price": round(price,2), "pred_units": float(max(0.0, min(pred_units, 1e7))), "elasticity": elasticity}
            agent_explanations = {}
            candidate_explanations = {}
            for aname, attr in traces.items():
                subj = {"agent": aname, "candidates": attr.candidates, "details": attr.details}
                try:
                    messages = build_explain_messages({"sku": row.get("sku_id")}, subj, {k: v.candidates for k,v in traces.items()}, sop)
                    agent_explanations[aname] = bedrock_invoke(messages, max_tokens=LLM_MAX_TOKENS, temperature=0.0) if boto3 else f"[fallback] Agent {aname} produced {attr.candidates}"
                except Exception as e:
                    agent_explanations[aname] = f"[fallback] Agent {aname} produced {attr.candidates} (LLM error: {e})"
            for key, info in candidates_detail.items():
                subj = {"candidate_key": key, "price": info["price"], "pred_units": info["pred_units"], "elasticity": info["elasticity"]}
                try:
                    messages = build_explain_messages({"sku": row.get("sku_id")}, subj, {k:v.candidates for k,v in traces.items()}, sop)
                    candidate_explanations[key] = bedrock_invoke(messages, max_tokens=LLM_MAX_TOKENS, temperature=0.0) if boto3 else f"[fallback] Candidate {key}: {info}"
                except Exception as e:
                    candidate_explanations[key] = f"[fallback] Candidate {key}: {info} (LLM error: {e})"
            selection = selector_and_sop(row, mapped, guardrails)
            try:
                final_expl = bedrock_invoke(build_explain_messages({"sku": row.get("sku_id")}, {"selection": selection}, {k:v.candidates for k,v in traces.items()}, sop), max_tokens=LLM_MAX_TOKENS, temperature=0.0) if boto3 else f"[fallback] selection {selection}"
            except Exception as e:
                final_expl = f"[fallback] Final selection {selection} (LLM error: {e})"
            print_agent_traces(traces)
            print_candidates_table(candidates_detail, candidate_explanations, selection)
            console.print(Panel("[bold]Agent explanations (short)[/bold]\n" + "\n\n".join(f"{k}: {v}" for k,v in agent_explanations.items())))
            console.print(Panel("[bold]Final explanation (short)[/bold]\n" + final_expl))
            out = {
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
                "sku": row.get("sku_id"),
                "channel": row.get("channel"),
                "region": row.get("region"),
                "agents": {k: v.candidates for k,v in traces.items()},
                "candidates": candidates_detail,
                "selection": selection,
                "agent_explanations": agent_explanations,
                "candidate_explanations": candidate_explanations,
                "final_explanation": final_expl,
                "sop_source": sop.get("source")
            }
            out_path = os.path.join(ensure_output_dir(), f"decision_{row.get('sku_id')}_{int(time.time())}.json")
            with open(out_path, "w") as f:
                json.dump(serialize_pipeline_result(out), f, indent=2)
            console.print(f"Decision logged to {out_path}", markup=False)

        elif ch == "2":
            sku = input("SKU_ID (for base): ").strip()
            channel = input("Channel: ").strip()
            region = input("Region: ").strip()
            sku = resolve_input_value(sku, sorted(df["sku_id"].astype(str).unique().tolist()), "SKU_ID")
            channel = resolve_input_value(channel, sorted(df["channel"].astype(str).unique().tolist()), "Channel")
            region = resolve_input_value(region, sorted(df["region"].astype(str).unique().tolist()), "Region")
            console.print("Enter overrides (leave blank to skip). Type 'nl' to enter a natural language scenario instead.")
            method = input("Type 'nl' for natural language or press Enter to input fields manually: ").strip().lower()
            overrides = {}
            nl_text = None
            if method == "nl":
                nl_text = input("Enter your natural-language scenario (e.g. 'If competitor drops to 95 and inventory is 8 days what happens?'): ").strip()
                bedrock_result = nl_interpret_with_bedrock(nl_text, {"sku": sku, "channel": channel, "region": region})
                # bedrock_result is expected to be dict with overrides key
                if isinstance(bedrock_result, dict) and "overrides" in bedrock_result:
                    parsed_overrides = bedrock_result.get("overrides", {})
                    console.print(Panel("[bold]Parsed Scenario (LLM+rules)[/bold]\n" + json.dumps({"parsed": parsed_overrides, "confidence": bedrock_result.get("extraction_confidence", "unknown"), "notes": bedrock_result.get("notes", "")}, indent=2)))
                    overrides.update(parsed_overrides)
                else:
                    # fallback to rule-based parse
                    parsed = parse_nl_scenario(nl_text)
                    console.print(Panel("[bold]Parsed Scenario (fallback parser)[/bold]\n" + json.dumps(parsed, indent=2)))
                    overrides.update(parsed)
            else:
                lp = input("Local Price override: ").strip()
                if lp:
                    overrides["local_price"] = float(lp)
                cp = input("Competitor Price override: ").strip()
                if cp:
                    overrides["competitor_price"] = float(cp)
                promo = input("Promo Depth % override: ").strip()
                if promo:
                    overrides["promo_depth_pct"] = float(promo)
                invd = input("Inventory Days override (gives inventory_days): ").strip()
                if invd:
                    overrides["inventory_days"] = float(invd)

            # build baseline row
            mask_exact = (df["sku_id"].astype(str) == sku)
            if mask_exact.sum() > 0:
                row = df[mask_exact].iloc[0]
            else:
                console.print("No exact match; using aggregated medians.", markup=False)
                row = pd.Series()
                row["sku_id"] = sku or "SYNTHETIC"
                row["channel"] = channel or df["channel"].mode().iloc[0]
                row["region"] = region or df["region"].mode().iloc[0]
                row["category"] = df["category"].mode().iloc[0]
                for col in ["local_price","competitor_price","promo_depth_pct","final_price","units_sold","inventory_level","revenue","margin","festival_lift","seasonality_index","lead_time_days","purchase_cost","mrp","tax_pct"]:
                    row[col] = float(df[col].median(skipna=True))
                row["date"] = df["date"].max()

            # pipeline function
            def pipeline(r):
                b = BasePriceAgent().run(r, df)
                p = PromoAgent(guardrails).run(r, b.candidates)
                c = CompetitorAgent().run(r, p.candidates)
                i = InventoryAgent().run(r, c.candidates)
                pr = ProcurementAgent().run(r, i.candidates)
                bl = BillingAgent().run(r, pr.candidates)
                traces = {"base": b, "promo": p, "competitor": c, "inventory": i, "procurement": pr, "billing": bl}
                elasticity = estimate_elasticity(df, sku=r.get("sku_id"), category=r.get("category"))
                for t in traces.values():
                    t.details["estimated_elasticity"] = elasticity
                mapped = CandidateAssembler(guardrails).assemble(traces)
                details = {}
                for k, pr in mapped.items():
                    pred_units = predict_units_for_candidate(r, pr, elasticity, nn_model, nn_scaler)
                    details[k] = {"price": round(pr, 2), "pred_units": float(max(0.0, min(pred_units, 1e7))), "elasticity": elasticity}
                sel = selector_and_sop(r, mapped, guardrails)
                return {"traces": traces, "candidates": details, "selection": sel}

            orig = pipeline(row)

            # apply overrides to a copy (coerce known keys)
            wrow = row.copy()
            # apply competitor_pct_change if provided (relative); compute absolute comp price if possible
            if "competitor_pct_change" in overrides and "competitor_price" not in overrides:
                base_comp = safe_float(row.get("competitor_price"), None)
                if base_comp is not None:
                    pct = overrides["competitor_pct_change"]
                    overrides["competitor_price"] = round(base_comp * (1 + pct/100.0), 2)
            for k, v in overrides.items():
                if k in ("local_price","competitor_price","promo_depth_pct","purchase_cost","mrp","tax_pct"):
                    try:
                        wrow[k] = float(v)
                    except Exception:
                        wrow[k] = v
                elif k in ("inventory_days","lead_time_days"):
                    wrow["inventory_days"] = float(v)
                    if k == "lead_time_days":
                        wrow["lead_time_days"] = float(v)
                elif k == "channel":
                    wrow["channel"] = resolve_input_value(str(v), sorted(df["channel"].astype(str).unique().tolist()), "Channel")
                elif k == "sku_id":
                    wrow["sku_id"] = resolve_input_value(str(v), sorted(df["sku_id"].astype(str).unique().tolist()), "SKU_ID")

            what = pipeline(wrow)

            # compute delta prices (guard against missing keys)
            delta = {}
            for k in orig["candidates"].keys():
                try:
                    delta[k] = what["candidates"][k]["price"] - orig["candidates"][k]["price"]
                except Exception:
                    delta[k] = None

            # create constrained LLM summary (LLM must only summarize engine outputs)
            nl_response = ""
            try:
                ctx = {
                    "sku": sku,
                    "overrides": overrides,
                    "original_selection": orig["selection"],
                    "what_if_selection": what["selection"],
                    "delta": delta
                }
                if boto3:
                    system = {
                        "role": "system",
                        "content": (
                            "You are a strict summarizer. Given the engine output JSON in 'context', produce a short 3-5 sentence plain-language summary. "
                            "Do NOT propose new prices or changes that contradict the engine output. If asked, produce only a short recommendation that matches the engine's selection and approval level."
                        )
                    }
                    user = {"role": "user", "content": json.dumps(ctx)}
                    nl_response = bedrock_invoke([system, user], max_tokens=200, temperature=0.0)
                    # Safety: force-check that LLM did not invent a different price by parsing numbers in its output.
                    # If LLM mentions a number that does not match engine selection, append a safety note.
                    selected_price = what["selection"].get("price_recommended")
                    price_mentions = re.findall(r"(?:₹|\b)([0-9]+(?:\.[0-9]+)?)", nl_response)
                    mismatched = False
                    for pm in price_mentions:
                        try:
                            if abs(float(pm) - float(selected_price)) > 1e-6 and float(pm) not in (safe_float(orig["selection"].get("price_recommended"), -9999),):
                                mismatched = True
                                break
                        except Exception:
                            continue
                    if mismatched:
                        nl_response += "\n\n[NOTE] LLM mentioned a price that does not match engine output; trust the engine's recommended price above."
                else:
                    nl_lines = []
                    nl_lines.append(f"What-if applied for SKU {sku}.")
                    nl_lines.append(f"Original recommended price: {orig['selection'].get('price_recommended')}, What-if recommended: {what['selection'].get('price_recommended')}.")
                    nl_lines.append(f"Delta (optimal): {delta.get('price_optimal', 0.0):+.2f}.")
                    nl_lines.append(f"Approval level: {what['selection'].get('approval')}.")
                    nl_response = "\n".join(nl_lines)
            except Exception as e:
                nl_response = f"[fallback] Could not generate NL summary (error: {e})"

            console.print("[bold]Original selection:[/bold]\n" + json.dumps(orig["selection"], indent=2))
            console.print("[bold]What-if selection:[/bold]\n" + json.dumps(what["selection"], indent=2))
            console.print("[bold]Delta (price changes):[/bold]\n" + json.dumps(delta, indent=2))
            console.print(Panel("[bold]Natural language summary (constrained):[/bold]\n" + nl_response))

            out_obj = {
                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S"),
                "sku": sku,
                "overrides": overrides,
                "original": serialize_pipeline_result(orig),
                "what_if": serialize_pipeline_result(what),
                "delta": delta,
                "nl_summary": nl_response
            }
            out_path = os.path.join(ensure_output_dir(), f"whatif_{sku}_{int(time.time())}.json")
            with open(out_path, "w") as f:
                json.dump(out_obj, f, indent=2)
            console.print(f"[green]What-if saved to: {out_path}[/green]", markup=False)

        elif ch == "3":
            console.print("[bold]Batch run over first 100 rows...[/bold]")
            rows = df.head(100)
            results = []
            for _, r in rows.iterrows():
                b = BasePriceAgent().run(r, df)
                p = PromoAgent(guardrails).run(r, b.candidates)
                c = CompetitorAgent().run(r, p.candidates)
                i = InventoryAgent().run(r, c.candidates)
                pr = ProcurementAgent().run(r, i.candidates)
                bl = BillingAgent().run(r, pr.candidates)
                traces = {"base": b, "promo": p, "competitor": c, "inventory": i, "procurement": pr, "billing": bl}
                mapped = CandidateAssembler(guardrails).assemble(traces)
                sel = selector_and_sop(r, mapped, guardrails)
                results.append({"sku": r.get("sku_id"), **mapped, "final": sel})
            out_csv = os.path.join(ensure_output_dir(), "batch_output.csv")
            pd.DataFrame(results).to_csv(out_csv, index=False)
            console.print(f"Batch saved to {out_csv}", markup=False)

        elif ch == "4":
            if not ENABLE_NN:
                console.print("NN disabled or torch not installed.", markup=False)
                continue
            console.print("[bold]Training NN demand model (this may take several minutes)...[/bold]")
            try:
                train_nn_demand(DATA_CSV, model_out="nn_demand_model.pt", scaler_out="nn_scaler.npz", epochs=8)
            except Exception as e:
                console.print(f"NN training failed: {e}", markup=False)

        elif ch == "5":
            if not ENABLE_RL:
                console.print("RL disabled or sklearn not installed.", markup=False)
                continue
            console.print("[bold]Training policy (RandomForest) ...[/bold]")
            try:
                train_policy_simple(DATA_CSV, policy_out="policy_rf.pkl", sample_size=30000)
            except Exception as e:
                console.print(f"Policy training failed: {e}", markup=False)

        elif ch == "6":
            sku = input("SKU_ID to forecast: ").strip()
            channel = input("Channel: ").strip()
            region = input("Region: ").strip()
            sku = resolve_input_value(sku, sorted(df["sku_id"].astype(str).unique().tolist()), "SKU_ID")
            channel = resolve_input_value(channel, sorted(df["channel"].astype(str).unique().tolist()), "Channel")
            region = resolve_input_value(region, sorted(df["region"].astype(str).unique().tolist()), "Region")
            start_date = input("Forecast start date (YYYY-MM-DD) [default tomorrow]: ").strip()
            if not start_date:
                start_date = (pd.Timestamp.now() + pd.Timedelta(days=1)).strftime("%Y-%m-%d")
            days = input("Days to forecast (default 30): ").strip()
            days = int(days) if days else 30
            console.print("[bold]Running forecast & pricing for the requested future period...[/bold]")
            res = forecast_and_price(df, sop, sku, channel, region, start_date, days, nn_model, nn_scaler, policy_model)
            out_path = os.path.join(ensure_output_dir(), f"forecast_price_{sku}_{start_date}.json")
            with open(out_path, "w") as f:
                json.dump(res, f, indent=2)
            console.print(f"Forecasted pricing saved: {out_path}", markup=False)
            for d in res[:min(5, len(res))]:
                sel = d.get("selection", {})
                cands = d.get("candidates", {})
                cand_str = ", ".join([f"{k}:{format_inr(v['price'])} (units:{v['pred_units']:.0f})" for k,v in cands.items()])
                console.print(Panel(f"Date: {d['date']}  Seasonality: {d['seasonality_index']}\nTrend(monthly%): {d.get('sku_trend_monthly_pct'):.4f}\nSelection: {json.dumps(sel)}\nCandidates: {cand_str}"))

        else:
            console.print("Invalid option", markup=False)

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="AI Sahayak engine")
    parser.add_argument("--mode", choices=["cli", "api"], default=os.getenv("AI_SAHAYAK_MODE", "cli"))
    parser.add_argument("--host", default=os.getenv("AI_SAHAYAK_API_HOST", "127.0.0.1"))  # Use 0.0.0.0 on EC2 so Lambda can reach /api
    parser.add_argument("--port", type=int, default=int(os.getenv("AI_SAHAYAK_API_PORT", "8000")))
    parser.add_argument("--reload-runtime", action="store_true", help="Force reload dataset/models at API start")
    args = parser.parse_args()
    try:
        if args.mode == "api":
            if args.reload_runtime:
                init_runtime(force_reload=True)
            api_app = create_api_app()
            console.print(f"Starting API at http://{args.host}:{args.port}", markup=False)
            rules = [r.rule for r in api_app.url_map.iter_rules() if "drive-ui" in r.rule or "chat-action" in r.rule]
            if rules:
                console.print(f"Chat drive-ui routes: {rules}", markup=False)
            api_app.run(host=args.host, port=args.port, debug=False)
        else:
            run_interactive()
    except Exception as e:
        console.print(f"Fatal error: {e}", markup=False)
        raise
